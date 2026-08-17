"""Build the mid-semester capstone technical report as a formatted .docx.

Formatting follows the Report Format Guidelines supplied by the department:
A4 portrait; margins 1" top/bottom/right and 1.5" left; Times New Roman;
16 pt bold chapter names, 14 pt bold headings, 12 pt body; 1.5 line spacing;
page numbers bottom centre; table content and captions at 10 pt; IEEE numeric
references, single spaced and justified.

Run:  python report/build_report.py
"""
from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = pathlib.Path(__file__).parent
FIG = HERE / "figures"
OUT = HERE / "DeepErase_MidSemester_Report.docx"

FONT = "Times New Roman"
BODY_PT = 12
HEAD_PT = 14
CHAP_PT = 16
CAP_PT = 10
TBL_PT = 10
LINE = 1.5

RED = RGBColor(0xC0, 0x00, 0x00)


# ---------------------------------------------------------------------------
# low-level helpers
# ---------------------------------------------------------------------------
def _set_font(run, size=BODY_PT, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), FONT)


def para(doc, text="", size=BODY_PT, bold=False, italic=False, align=None,
         line=LINE, space_after=6, space_before=0, indent=None, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        _set_font(p.add_run(text), size, bold, italic, color)
    return p


def body(doc, text, **kw):
    return para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, **kw)


def bullet(doc, text, size=BODY_PT, indent=0.35):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = LINE
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(indent)
    _set_font(p.add_run(text), size)
    return p


def rule(doc):
    """The heavy horizontal rule under each chapter/front-matter title."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.space_before = Pt(0)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    p._element.get_or_add_pPr().append(pbdr)
    return p


def chapter(doc, title, page_break=True):
    if page_break:
        doc.add_page_break()
    para(doc, title.upper(), size=CHAP_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    rule(doc)


def heading(doc, text):
    return para(doc, text, size=HEAD_PT, bold=True, space_before=10,
                space_after=4)


def subheading(doc, text):
    return para(doc, text, size=BODY_PT, bold=True, space_before=8,
                space_after=3)


def caption(doc, text, above=False):
    return para(doc, text, size=CAP_PT, align=WD_ALIGN_PARAGRAPH.CENTER,
                line=1.0, space_after=10 if not above else 3,
                space_before=3 if not above else 8)


def figure(doc, name, number, text, width=6.0):
    path = FIG / f"{name}.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, f"FIGURE {number}: {text}")


def table(doc, number, title, header, rows, widths=None, font=TBL_PT):
    if number != "":
        caption(doc, f"TABLE {number}: {title}", above=True)
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(1)
        _set_font(p.add_run(h), font, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(1)
            _set_font(p.add_run(str(val)), font)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    para(doc, "", space_after=4)
    return t


def _field(paragraph, instr):
    r = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = instr
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r._r.append(f1); r._r.append(it); r._r.append(f2)
    _set_font(r, BODY_PT)


def page_numbers(section, fmt="decimal", start=None):
    """Bottom-centre page number, with an explicit numbering format."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    _field(p, " PAGE ")

    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def setup_section(section):
    section.page_width = Inches(8.27)      # A4 portrait
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.left_margin = Inches(1.5)


# ---------------------------------------------------------------------------
# document
# ---------------------------------------------------------------------------
def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(BODY_PT)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.line_spacing = LINE

    setup_section(doc.sections[0])

    # ---------------- cover page -----------------------------------------
    para(doc, "", space_after=24)
    para(doc, "DEEPERASE: MEASURING THE DEPTH AND BREADTH OF KNOWLEDGE "
              "REMOVAL IN LARGE LANGUAGE MODEL UNLEARNING",
         size=CHAP_PT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=18)
    for t, s, b in [("Capstone Project Report", HEAD_PT, True),
                    ("MID SEMESTER EVALUATION", HEAD_PT, True),
                    ("Submitted by:", BODY_PT, True)]:
        para(doc, t, size=s, bold=b, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=10)
    for _ in range(5):
        para(doc, "(<University Roll Number>) NAME OF THE STUDENT",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
             color=RED)
    for t in ["BE Third Year, CoE/CoSE", "CPG No: 221", "Under the Mentorship of",
              "Dr. Suresh Kumar Chaudhary", "Assistant Professor"]:
        para(doc, t, bold=("BE Third" in t or "CPG" in t),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, "", space_after=18)
    para(doc, "Computer Science and Engineering Department", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Thapar Institute of Engineering and Technology, Patiala",
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "August 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---------------- front matter (roman numerals) ----------------------
    front = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(front)
    page_numbers(front, fmt="lowerRoman", start=2)

    # ABSTRACT
    para(doc, "ABSTRACT", size=CHAP_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    rule(doc)
    body(doc,
         "Large language models absorb information during training that may "
         "later have to be removed, whether because a person withdraws consent, "
         "because copyrighted material was included, or because a capability "
         "turns out to be unsafe. Retraining the model from scratch without the "
         "data is the only method guaranteed to work, and it is far too "
         "expensive to use routinely. Machine unlearning aims to achieve the "
         "same result by modifying an already-trained model. Whether it "
         "genuinely does so is unsettled: a model can be made to stop producing "
         "an answer while still holding the underlying knowledge internally, and "
         "most current evaluation cannot tell the two situations apart.")
    body(doc,
         "This project builds a measurement system that separates two distinct "
         "properties of an unlearned model, which the literature normally "
         "conflates. Depth asks whether the knowledge is still physically "
         "present in the network's internal representations; it is measured by "
         "two-stage activation patching, following the Unlearning Depth Score. "
         "Breadth asks how far removal generalises beyond the exact wording the "
         "model was trained on; it is measured by tiered forced-choice probes. "
         "Both axes are calibrated against the same pair of reference models, so "
         "that 0 always means the original model and 1 always means a model "
         "retrained without the data, and the two numbers can be read on one "
         "scale.")
    body(doc,
         "The depth axis reproduces published reference values to a mean "
         "absolute deviation of 0.010, against the 0.08 agreement threshold we set in advance, across "
         "seven runs and two model scales. The breadth axis discriminates "
         "correctly between a model that learned the target facts and one that "
         "never saw them, and is calibrated to that measured range rather than "
         "reported on a nominal zero-to-one scale. Using these instruments the "
         "system produces a complete depth-breadth trajectory: a model is "
         "unlearned by gradient ascent with checkpoint selection, and the "
         "resulting update is then scaled by an extrapolation parameter to trace "
         "a controlled path through the plane. The implementation is a tested "
         "Python package with 489 automated tests.")
    subheading(doc, "Keywords")
    body(doc, "machine unlearning; large language models; activation patching; "
              "mechanistic interpretability; TOFU benchmark; model editing; "
              "evaluation methodology.", space_after=0)

    # DECLARATION
    doc.add_page_break()
    para(doc, "DECLARATION", size=CHAP_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    rule(doc)
    body(doc,
         "We hereby declare that the design principles and working prototype "
         "model of the project entitled DeepErase: Measuring the Depth and "
         "Breadth of Knowledge Removal in Large Language Model Unlearning is an "
         "authentic record of our own work carried out in the Computer Science "
         "and Engineering Department, TIET, Patiala, under the guidance of "
         "Dr. Suresh Kumar Chaudhary during the 5th semester (2026).")
    para(doc, "Date:", space_before=10, space_after=10)
    table(doc, "", "", ["Roll No.", "Name", "Signature"],
          [["", "", "----"] for _ in range(5)], widths=[1.6, 2.6, 1.4])
    para(doc, "Counter Signed By:", italic=True, space_before=14, space_after=14)
    para(doc, "Faculty Mentor:", space_after=18)
    para(doc, "Dr. Suresh Kumar Chaudhary", space_after=2, indent=0.5)
    para(doc, "Assistant Professor", space_after=2, indent=0.5)
    para(doc, "CSED,", space_after=2, indent=0.5)
    para(doc, "TIET, Patiala", space_after=2, indent=0.5)

    # ACKNOWLEDGEMENT
    doc.add_page_break()
    para(doc, "ACKNOWLEDGEMENT", size=CHAP_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    rule(doc)
    body(doc,
         "We would like to express our thanks to our mentor, "
         "Dr. Suresh Kumar Chaudhary. He has been of great help in our venture "
         "and an indispensable resource of technical knowledge, and his "
         "insistence on measuring what we claim rather than what we hoped for "
         "shaped the direction of this work.")
    body(doc,
         "We are also thankful to the Head, Computer Science and Engineering "
         "Department, the entire faculty and staff of the Computer Science and "
         "Engineering Department, and also our friends who devoted their "
         "valuable time and helped us in all possible ways towards the "
         "successful completion of this project. We thank all those who have "
         "contributed either directly or indirectly towards this project.")
    body(doc,
         "Lastly, we would also like to thank our families for their unyielding "
         "love and encouragement. They always wanted the best for us and we "
         "admire their determination and sacrifice.")
    para(doc, "Date:", space_before=10, space_after=10)
    table(doc, "", "", ["Roll No.", "Name", "Signature"],
          [["", "", "----"] for _ in range(5)], widths=[1.6, 2.6, 1.4])

    # TABLE OF CONTENTS
    doc.add_page_break()
    para(doc, "TABLE OF CONTENTS", size=CHAP_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    rule(doc)
    toc = [
        ("ABSTRACT", "ii", True), ("DECLARATION", "iii", True),
        ("ACKNOWLEDGEMENT", "iv", True), ("LIST OF FIGURES", "vii", True),
        ("LIST OF TABLES", "viii", True), ("LIST OF ABBREVIATIONS", "ix", True),
        ("", "", False),
        ("CHAPTER 1  INTRODUCTION", "1", True),
        ("     1.1  Project Overview", "1", False),
        ("     1.2  Need Analysis", "4", False),
        ("     1.3  Research Gaps", "5", False),
        ("     1.4  Problem Definition and Scope", "7", False),
        ("     1.5  Assumptions and Constraints", "8", False),
        ("     1.6  Standards", "9", False),
        ("     1.7  Approved Objectives", "10", False),
        ("     1.8  Methodology", "11", False),
        ("     1.9  Project Outcomes and Deliverables", "12", False),
        ("     1.10 Novelty of Work", "13", False),
        ("CHAPTER 2  REQUIREMENT ANALYSIS", "14", True),
        ("     2.1  Literature Survey", "14", False),
        ("     2.2  Software Requirement Specification", "20", False),
        ("     2.3  Cost Analysis", "25", False),
        ("     2.4  Risk Analysis", "26", False),
        ("CHAPTER 3  METHODOLOGY ADOPTED", "27", True),
        ("     3.1  Investigative Techniques", "27", False),
        ("     3.2  Proposed Solution", "30", False),
        ("     3.3  Work Breakdown Structure", "33", False),
        ("     3.4  Tools and Technology", "34", False),
        ("CHAPTER 4  DESIGN SPECIFICATIONS", "35", True),
        ("     4.1  System Architecture", "35", False),
        ("     4.2  Design Level Diagrams", "37", False),
        ("     4.3  User Interface Diagrams", "41", False),
        ("     4.4  Snapshots of Working Prototype", "42", False),
        ("CHAPTER 5  CONCLUSIONS AND FUTURE SCOPE", "47", True),
        ("     5.1  Work Accomplished", "47", False),
        ("     5.2  Conclusions", "48", False),
        ("     5.3  Environmental, Economic and Social Benefits", "49", False),
        ("     5.4  Future Work Plan", "50", False),
        ("APPENDIX A: REFERENCES", "52", True),
        ("APPENDIX B: PLAGIARISM REPORT", "55", True),
    ]
    for label, page, bold in toc:
        if not label:
            para(doc, "", space_after=2)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = LINE
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(5.4), WD_ALIGN_PARAGRAPH.RIGHT)
        _set_font(p.add_run(f"{label}\t{page}"), BODY_PT, bold=bold)

    # LIST OF FIGURES
    doc.add_page_break()
    para(doc, "LIST OF FIGURES", size=CHAP_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    rule(doc)
    figs = [
        ("Figure 1", "Work breakdown structure and schedule", "33"),
        ("Figure 2", "System architecture of the DeepErase measurement pipeline", "35"),
        ("Figure 3", "Two-stage activation patching used to compute the depth score", "37"),
        ("Figure 4", "Breadth tier structure: one fact probed at increasing semantic distance", "38"),
        ("Figure 5", "Class diagram of the principal modules", "39"),
        ("Figure 6", "Sequence diagram for a complete depth-breadth study run", "40"),
        ("Figure 7", "State chart for the model lifecycle and checkpoint selection", "41"),
        ("Figure 8", "Use case diagram for the overall system", "42"),
        ("Figure 9", "Depth-metric validation against published reference values", "43"),
        ("Figure 10", "Measured unlearning trajectory and checkpoint selection", "45"),
        ("Figure 11", "Measured depth-breadth trajectory under extrapolation", "46"),
    ]
    _list_table(doc, ["Figure No.", "Caption", "Page No."], figs)

    # LIST OF TABLES
    doc.add_page_break()
    para(doc, "LIST OF TABLES", size=CHAP_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    rule(doc)
    tabs = [
        ("Table 1", "Research gaps identified from the literature", "5"),
        ("Table 2", "Assumptions and constraints", "8"),
        ("Table 3", "Approved objectives and current status", "10"),
        ("Table 4", "Literature survey summary", "17"),
        ("Table 5", "Cost analysis", "25"),
        ("Table 6", "Risk analysis", "26"),
        ("Table 7", "Investigative techniques considered", "27"),
        ("Table 8", "Work breakdown structure", "33"),
        ("Table 9", "Tools and technology", "34"),
        ("Table 10", "Depth-metric validation against published values", "42"),
        ("Table 11", "Breadth-axis calibration measurements", "43"),
        ("Table 12", "Measured depth-breadth trajectory", "45"),
        ("Table 13", "Automated test suite coverage", "46"),
    ]
    _list_table(doc, ["Table No.", "Caption", "Page No."], tabs)

    # LIST OF ABBREVIATIONS
    doc.add_page_break()
    para(doc, "LIST OF ABBREVIATIONS", size=CHAP_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    rule(doc)
    abbr = [
        ("LLM", "Large Language Model"),
        ("UDS", "Unlearning Depth Score"),
        ("KE layer", "Knowledge-Encoding layer"),
        ("LER", "Layer Erasure Ratio"),
        ("GA", "Gradient Ascent"),
        ("GradDiff", "Gradient Difference"),
        ("NPO", "Negative Preference Optimization"),
        ("UIPE", "Unlearning Improvement via Parameter Extrapolation"),
        ("TOFU", "Task of Fictitious Unlearning"),
        ("NLL", "Negative Log-Likelihood"),
        ("QA", "Question Answering"),
        ("LoRA", "Low-Rank Adaptation"),
        ("SRS", "Software Requirement Specification"),
        ("WBS", "Work Breakdown Structure"),
        ("API", "Application Programming Interface"),
        ("CLI", "Command Line Interface"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for a, full in abbr:
        cells = t.add_row().cells
        for i, val in enumerate((a, full)):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(1)
            _set_font(p.add_run(val), TBL_PT, bold=(i == 0))
        cells[0].width = Inches(1.4)
        cells[1].width = Inches(4.4)

    # ---------------- body (arabic numerals) ------------------------------
    main = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(main)
    page_numbers(main, fmt="decimal", start=1)

    _chapter1(doc)
    _chapter2(doc)
    _chapter3(doc)
    _chapter4(doc)
    _chapter5(doc)
    _appendices(doc)

    doc.save(OUT)
    print(f"Wrote {OUT}")
    return OUT


def _list_table(doc, header, rows):
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        _set_font(p.add_run(h), TBL_PT, bold=True)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(1)
            _set_font(p.add_run(val), TBL_PT)
        cells[0].width = Inches(1.0)
        cells[1].width = Inches(4.2)
        cells[2].width = Inches(0.8)


# ===========================================================================
# CHAPTER 1
# ===========================================================================
def _chapter1(doc):
    chapter(doc, "Chapter 1  Introduction", page_break=False)

    heading(doc, "1.1  Project Overview")
    body(doc,
         "A large language model is trained by adjusting several billion "
         "numerical parameters so that the model becomes good at predicting the "
         "next word in its training text. Nothing in that process keeps the "
         "training documents in a form that can be looked up or deleted later. "
         "The information is distributed across the parameters, and once "
         "training is finished there is no index, no record of which document "
         "contributed to which parameter, and no delete operation.")
    body(doc,
         "This becomes a problem the moment some of that information must be "
         "removed. A person may exercise a right to have their personal data "
         "erased. A rights holder may establish that copyrighted text was used "
         "without permission. An operator may discover that the model has "
         "acquired a capability that is unsafe to expose. In each case the "
         "obligation is to remove specific knowledge from a model that is "
         "already trained and already deployed.")
    body(doc,
         "The one method that is certain to work is to delete the offending "
         "data from the training corpus and train the model again from the "
         "beginning. The resulting model provably never saw the data. It is "
         "also, for a model of any realistic size, prohibitively expensive in "
         "time, money and energy, and it must be repeated for every subsequent "
         "removal request. Machine unlearning is the study of methods that aim "
         "to reach the same end state by modifying the parameters of the "
         "existing model instead [18].")
    body(doc,
         "The central difficulty is verification. When an unlearning method has "
         "been applied, the natural way to check it is to ask the model the "
         "question and see whether the correct answer still comes out. If it "
         "does not, the method appears to have worked. That test, however, "
         "cannot distinguish between two very different outcomes. In the first, "
         "the knowledge has genuinely been removed from the parameters. In the "
         "second, the knowledge is still encoded exactly as before and the model "
         "has merely learned to suppress it at the output. The second outcome "
         "looks identical to the first under question-and-answer testing, but it "
         "is not a removal at all; the information remains recoverable, and a "
         "growing body of work shows that it can in fact be recovered.")
    body(doc,
         "Several results establish this concretely. Lynch et al. [11] show that "
         "models that pass standard unlearning evaluations still reveal the "
         "supposedly removed content under alternative prompting. Łucki et al. "
         "[12] demonstrate recovery under adversarial pressure and argue that "
         "current evaluations systematically overstate how much has been "
         "removed. Deeb and Roger [13] show that a small amount of retraining on "
         "related data can bring back knowledge that unlearning appeared to have "
         "destroyed, which is only possible if the knowledge was still present. "
         "Doshi and Stickland [14] reach a similar conclusion from purely "
         "black-box probing. The consistent finding is that suppression is "
         "common, that it is easily mistaken for removal, and that the standard "
         "evaluation instruments cannot separate the two.")
    body(doc,
         "This project treats that separation as the object of study rather than "
         "as an obstacle. It builds a measurement system that reports two "
         "distinct properties of an unlearned model.")
    body(doc,
         "The first property is depth. Depth asks whether the knowledge is still "
         "physically encoded inside the network, independently of what the model "
         "chooses to say. It is measured mechanistically, by activation "
         "patching: internal activations are copied from one model into another "
         "at a specific layer and a specific token position, and the effect on "
         "the model's predicted probability for the target fact is recorded. If "
         "copying activations out of an unlearned model into a model that still "
         "knows the fact fails to disturb that fact, the unlearned model is no "
         "longer carrying it. If it disturbs it just as much as before, the "
         "knowledge is intact and only the output behaviour changed. This "
         "project follows the Unlearning Depth Score formulation of Lee et al. "
         "[2], which turns that comparison into a single number anchored at 0 "
         "for an unchanged model and 1 for a model retrained without the data.")
    body(doc,
         "The second property is breadth. Breadth asks how far removal "
         "generalises. Unlearning is usually optimised against a specific set of "
         "training sentences, and a method can succeed on exactly those "
         "sentences while leaving the same fact fully accessible through a "
         "paraphrase, a different name for the same entity, or a question whose "
         "answer follows from the removed fact. Breadth is therefore measured "
         "behaviourally, using probes arranged in tiers of increasing distance "
         "from the training form, together with a control tier of knowledge that "
         "must not change at all.")
    body(doc,
         "The two properties are logically independent, and the relationship "
         "between them is the research question this project exists to answer. "
         "A method might remove a fact very thoroughly in exactly one phrasing "
         "and nowhere else; that is deep but narrow. A method might blunt the "
         "model's access to a fact across every phrasing without removing the "
         "underlying representation anywhere; that is broad but shallow. Both "
         "would be described simply as 'unlearning' by current practice, and "
         "both would receive similar scores on standard benchmarks, yet they "
         "carry entirely different privacy guarantees. Whether these two ways of "
         "forgetting trade off against each other, and how, is not established "
         "in the literature.")
    body(doc,
         "The experimental setting is the TOFU benchmark [1], which was designed "
         "for exactly this kind of study. TOFU consists of two hundred entirely "
         "fictitious author biographies, each with twenty question-answer pairs. "
         "Because the authors are invented, no model has encountered them during "
         "general pretraining, so any knowledge the model has of them came from "
         "a controlled fine-tuning step and can be reasoned about precisely. "
         "Critically, TOFU also publishes models that were retrained from the "
         "start while omitting a defined fraction of the authors. These "
         "retrained models are ground truth: they show what a perfectly "
         "unlearned model would look like. Every measurement in this project is "
         "calibrated against them, so that both axes read 0 for the original "
         "model and 1 for the retrained oracle, and a number can be interpreted "
         "as a fraction of the distance to genuine removal.")
    body(doc,
         "The deliverable at this stage is a working, tested measurement system "
         "together with evidence that it measures what it claims. The depth axis "
         "has been validated by reproducing published reference values. The "
         "breadth axis has been calibrated against models that definitely do and "
         "definitely do not hold the target knowledge. The complete pipeline has "
         "been run end to end to produce a first depth-breadth trajectory. "
         "Chapter 4 presents these results and Chapter 5 sets out what remains.")

    heading(doc, "1.2  Need Analysis")
    body(doc,
         "The need for unlearning is established by regulation, by litigation "
         "and by safety practice. Data protection regimes including the GDPR "
         "grant individuals a right to erasure, and that right does not "
         "obviously stop at a database boundary if personal information has been "
         "absorbed into a deployed model. Copyright disputes over training "
         "corpora raise the same question for rights holders. Safety reviews "
         "raise it for hazardous capability, which is the motivation behind "
         "benchmarks such as WMDP [8].")
    body(doc,
         "Retraining answers all of these correctly and is not a practical "
         "answer. Training a frontier model consumes months of compute; erasure "
         "requests arrive continuously. An organisation that could only comply "
         "by retraining would be unable to comply at all. Unlearning methods "
         "exist to close that gap, and a substantial family of them now does: "
         "gradient ascent and its variants [5], preference-optimisation "
         "objectives [4], representation-level interventions [8], and parameter "
         "extrapolation [3].")
    body(doc,
         "The unmet need is not another method. It is trustworthy measurement. "
         "An organisation that deletes personal data on the strength of an "
         "unlearning procedure is making a legal and ethical claim about that "
         "model. If the evaluation supporting the claim cannot distinguish "
         "removal from suppression, the claim is unfounded, and the evidence "
         "reviewed in Section 1.1 indicates that this is the normal situation "
         "rather than an edge case. Barez et al. [19] and Thaker et al. [20] "
         "both identify evaluation, not method design, as the field's binding "
         "constraint.")
    body(doc,
         "This project addresses that constraint directly. It provides "
         "instruments that report how deeply knowledge was removed and how "
         "broadly, on a scale anchored to retraining, so that the strength of an "
         "unlearning claim can be stated rather than assumed.")

    heading(doc, "1.3  Research Gaps")
    body(doc,
         "Six gaps were identified from the literature survey in Section 2.1. "
         "They are summarised in Table 1 and discussed below.")
    table(doc, 1, "Research gaps identified from the literature",
          ["No.", "Research gap", "Key references"],
          [["G1", "Evaluation is behavioural rather than mechanistic, so "
                  "output suppression is scored identically to representational "
                  "removal.", "[11], [12], [13], [14]"],
           ["G2", "Depth of erasure has only recently been formalised and has "
                  "never been related to how broadly a fact is forgotten.",
            "[2], [15]"],
           ["G3", "Breadth is under-specified. Benchmarks probe at most "
                  "paraphrase level, leaving alias, entailment and multi-hop "
                  "access unmeasured.", "[1], [9], [10]"],
           ["G4", "Parameter extrapolation is tuned for utility retention; its "
                  "effect on internal representations is uncharacterised.",
            "[3], [17]"],
           ["G5", "Scores are reported on a nominal zero-to-one scale although "
                  "the achievable range is much narrower, which inflates "
                  "apparent effect sizes.", "[1], [9]"],
           ["G6", "Results are sensitive to annotation and sampling choices that "
                  "are not reported, limiting reproducibility.", "[2], [16]"]],
          widths=[0.5, 3.7, 1.5])

    subheading(doc, "G1  Behavioural evaluation cannot see representational removal")
    body(doc,
         "Standard unlearning evaluation measures what the model outputs: "
         "answer accuracy, ROUGE overlap, membership-inference signal. All of "
         "these are functions of the model's behaviour. A model that has learned "
         "to withhold an answer scores the same as a model from which the answer "
         "has been removed. Lynch et al. [11] and Łucki et al. [12] both "
         "demonstrate the resulting failure directly, and Deeb and Roger [13] "
         "show that the information is often still present in the weights. The "
         "gap is the absence of a routine, quantitative measurement of "
         "representational removal.")

    subheading(doc, "G2  Depth has not been related to breadth")
    body(doc,
         "Lee et al. [2] provide the first practical scalar measure of how "
         "deeply knowledge has been erased, and Hong et al. [15] provide a "
         "complementary parametric view. Both treat depth in isolation. Neither "
         "asks whether a method that erases deeply also erases widely, or "
         "whether the two are in tension. That relationship is the specific "
         "opening this project targets.")

    subheading(doc, "G3  Breadth is measured only at paraphrase level")
    body(doc,
         "TOFU [1] supplies a paraphrased variant of each answer and a set of "
         "perturbed alternatives; MUSE [9] and RWKU [10] add further behavioural "
         "probes. None systematically probes an alias for the same entity, a "
         "fact entailed by the removed fact, or a two-step inference over it. "
         "Since these are precisely the routes by which suppressed knowledge "
         "resurfaces, the omission matters.")

    subheading(doc, "G4  Extrapolation is tuned for utility, not characterised for erasure")
    body(doc,
         "UIPE [3] improves unlearning by extrapolating along the update "
         "direction, an idea whose theoretical home is task arithmetic [17]. The "
         "extrapolation coefficient is selected to preserve general model "
         "utility. What the coefficient does to the internal encoding of the "
         "target knowledge is not reported. This project uses that coefficient "
         "as a controlled experimental variable rather than a hyperparameter to "
         "be tuned away.")

    subheading(doc, "G5  Uncalibrated scales inflate apparent effects")
    body(doc,
         "Forced-choice and accuracy metrics are conventionally reported as if "
         "0 and 1 were attainable endpoints. In practice a model that has never "
         "seen a fact still scores well above chance on such probes, because "
         "many distractors are implausible on general grounds. The usable range "
         "is therefore much narrower than the nominal range, and an improvement "
         "that looks small on the nominal scale may be large on the real one, or "
         "the reverse. Calibrating against reference models is the remedy this "
         "project adopts.")

    subheading(doc, "G6  Undocumented choices control the result")
    body(doc,
         "Depth measurement requires a decision about which tokens carry the "
         "target fact. Different reasonable policies produce materially "
         "different scores. Similarly, benchmark splits that are nested rather "
         "than disjoint make naive sampling unrepresentative. Neither choice is "
         "usually reported. This project treats both as explicit, "
         "version-controlled configuration.")

    heading(doc, "1.4  Problem Definition and Scope")
    subheading(doc, "Problem statement")
    body(doc,
         "Given a language model that has been subjected to an unlearning "
         "procedure, determine (i) how deeply the target knowledge has been "
         "removed from its internal representations, (ii) how broadly that "
         "removal generalises beyond the exact training phrasing, and (iii) how "
         "these two quantities relate to one another as the strength of the "
         "unlearning intervention is varied.")
    subheading(doc, "In scope")
    for s in [
        "Depth measurement by two-stage activation patching, following the "
        "published Unlearning Depth Score formulation, validated against its "
        "published reference values.",
        "Breadth measurement by tiered forced-choice probing at exact and "
        "paraphrase level, with a retained-knowledge control tier, calibrated "
        "against reference models.",
        "Unlearning by gradient ascent, gradient difference and negative "
        "preference optimisation, with checkpoint selection against a utility "
        "floor.",
        "Controlled traversal of the depth-breadth plane by parameter "
        "extrapolation.",
        "The TOFU benchmark and the TOFU-tuned Llama-3.2 model family at the 1B "
        "and 3B scales.",
    ]:
        bullet(doc, s)
    subheading(doc, "Out of scope")
    for s in [
        "Proposing a new unlearning method. This project measures existing "
        "methods; it does not compete with them.",
        "Models beyond the 3B scale, and closed-weight models, since both depth "
        "measurement and unlearning require full parameter access.",
        "Adversarial attack development. Robustness is measured, not attacked.",
        "Formal privacy guarantees such as differential privacy.",
    ]:
        bullet(doc, s)

    heading(doc, "1.5  Assumptions and Constraints")
    body(doc,
         "Table 2 records the assumptions on which the measurements rest and "
         "the constraints that bound what can be studied. The distinction "
         "matters: an assumption is something taken to be true that could in "
         "principle be wrong and would invalidate a result if it were, whereas "
         "a constraint limits scope without threatening validity. Assumptions 1 "
         "and 3 are the load-bearing ones, since the entire depth measurement "
         "depends on the retained models being a valid oracle and on entity "
         "representations being localisable to the positions that are patched.")
    table(doc, 2, "Assumptions and constraints",
          ["No.", "Type", "Statement"],
          [["1", "Assumption",
            "The TOFU retain models are a valid ground truth for perfect "
            "unlearning. They were retrained from the same initialisation with "
            "the forget authors excluded, so they never encountered the target "
            "knowledge."],
           ["2", "Assumption",
            "The fictitious authors in TOFU are absent from general pretraining "
            "corpora, so all measured knowledge of them originates in the "
            "controlled fine-tuning step."],
           ["3", "Assumption",
            "Hidden-state activations at the token positions immediately "
            "preceding an entity mention carry the model's encoding of that "
            "entity, which is the assumption underlying activation patching and "
            "is inherited from the causal-tracing literature [7]."],
           ["4", "Assumption",
            "Forced-choice ranking against perturbed alternatives is a valid "
            "proxy for whether a model retains a fact, given that the "
            "alternatives are generated from the same paraphrase as the correct "
            "answer."],
           ["5", "Constraint",
            "Depth measurement requires access to all model weights and to "
            "internal activations, restricting the study to open-weight models."],
           ["6", "Constraint",
            "Full-parameter unlearning requires memory for weights, gradients "
            "and two optimiser moments, approximately four times the memory of "
            "the weights alone. This bounds the model scale that can be studied."],
           ["7", "Constraint",
            "Breadth resolution is limited by the number of probe items: with N "
            "items the smallest observable change is one item in N, divided by "
            "the calibrated range."],
           ["8", "Constraint",
            "Gradient ascent degrades general model utility if run too long, so "
            "the usable range of unlearning strength is bounded by a utility "
            "floor rather than chosen freely."]],
          widths=[0.5, 1.0, 4.2])

    heading(doc, "1.6  Standards")
    for s in [
        "IEEE citation and referencing style for all bibliographic material, as "
        "specified in the departmental guidelines.",
        "PEP 8 and PEP 257 for Python source layout and documentation strings; "
        "type annotations throughout the package.",
        "Semantic versioning for the measurement package, so that a reported "
        "result can be tied to the exact implementation that produced it.",
        "Reproducibility practice consistent with ACM artefact review: every run "
        "writes a self-describing directory containing its configuration, its "
        "raw measurements and its environment.",
        "Licence compliance: the TOFU dataset and the derived model checkpoints "
        "are used under their published terms, and the Llama 3.2 community "
        "licence governs the base models.",
        "Ethical scope: all experimental data concerns fictitious individuals "
        "invented for benchmarking, so no personal data is processed at any "
        "point in this work.",
    ]:
        bullet(doc, s)

    heading(doc, "1.7  Approved Objectives")
    body(doc,
         "The objectives approved at the proposal evaluation are listed in "
         "Table 3 together with their present status.")
    table(doc, 3, "Approved objectives and current status",
          ["No.", "Objective", "Status"],
          [["O1", "Establish a reproducible experimental environment and data "
                  "pipeline over the TOFU benchmark and the TOFU-tuned model "
                  "family.", "Complete"],
           ["O2", "Implement a mechanistic depth measurement based on two-stage "
                  "activation patching, and validate it against published "
                  "reference values.", "Complete"],
           ["O3", "Design and implement a breadth measurement that probes "
                  "generalisation of forgetting beyond the training phrasing, "
                  "calibrated against reference models.", "Complete for tiers "
                                                          "B0, B1 and R"],
           ["O4", "Implement representative unlearning methods with a principled "
                  "stopping rule that prevents model collapse.", "Complete"],
           ["O5", "Traverse the depth-breadth plane under controlled parameter "
                  "extrapolation and characterise the relationship between the "
                  "two axes.", "First trajectory obtained; replication in "
                               "progress"],
           ["O6", "Report the relationship across multiple unlearning methods, "
                  "seeds and model scales.", "In progress"]],
          widths=[0.5, 3.9, 1.3])

    heading(doc, "1.8  Methodology")
    body(doc,
         "The work is organised as an experimental study with explicit controls, "
         "justified in detail in Section 3.1. The method has five stages.")
    for s in [
        "Anchor. Two reference models are measured: the model that learned the "
        "target facts, and the model retrained without them. These fix the two "
        "ends of both scales.",
        "Calibrate depth. A first patching pass from the retrained model into "
        "the original model establishes, for each example and each layer, how "
        "much of the fact that layer carries. Layers carrying a negligible share "
        "are discarded.",
        "Unlearn. An unlearning method is applied to the original model. Utility "
        "on retained knowledge is evaluated after every epoch, and the "
        "checkpoint returned is the one that forgot most while remaining above a "
        "utility floor.",
        "Traverse. The difference between the unlearned and original parameters "
        "defines an update direction. Scaling it by a coefficient produces a "
        "family of models of increasing unlearning strength along one controlled "
        "path.",
        "Measure. Every model on that path is scored for depth, for breadth, and "
        "for retained utility, the last acting as a control that distinguishes "
        "targeted forgetting from general damage.",
    ]:
        bullet(doc, s)

    heading(doc, "1.9  Project Outcomes and Deliverables")
    for s in [
        "A tested Python package implementing both measurement axes, three "
        "unlearning methods and the extrapolation traversal, with 489 automated "
        "tests.",
        "A validated depth measurement reproducing published reference values to "
        "a mean absolute deviation of 0.010, against a pre-set agreement threshold of 0.08.",
        "A calibrated breadth measurement with empirically determined floor and "
        "ceiling rather than assumed endpoints.",
        "Command-line tools for environment verification, depth validation, "
        "breadth measurement and the full study, each writing a self-describing "
        "run directory.",
        "A first complete depth-breadth trajectory, reported in Section 4.4.",
        "This technical report, and a final report characterising the "
        "depth-breadth relationship across methods and seeds.",
    ]:
        bullet(doc, s)

    heading(doc, "1.10  Novelty of Work")
    body(doc,
         "The novelty is not a new unlearning algorithm. It is the joint, "
         "calibrated measurement of two properties that the literature currently "
         "treats separately or not at all.")
    for s in [
        "Depth and breadth are measured on a single common scale, both anchored "
        "at 0 for the original model and 1 for the retrained oracle, so that the "
        "two can be compared directly. No prior work reports both.",
        "A mechanistic measure and a behavioural measure are applied to the same "
        "models under the same conditions, which is what makes the distinction "
        "between removal and suppression observable rather than assumed.",
        "The extrapolation coefficient is repurposed from a hyperparameter into "
        "a controlled experimental variable, producing a continuous trajectory "
        "through the plane instead of a single unlearned point.",
        "Breadth is reported against a measured dynamic range rather than a "
        "nominal one, and analysis is gated on whether that range is large "
        "enough to support a conclusion.",
    ]:
        bullet(doc, s)


# ===========================================================================
# CHAPTER 2
# ===========================================================================
def _chapter2(doc):
    chapter(doc, "Chapter 2  Requirement Analysis")

    heading(doc, "2.1  Literature Survey")

    subheading(doc, "2.1.1  Theory Associated With Problem Area")
    body(doc,
         "Machine unlearning originates in classical machine learning, where "
         "the objective is stated exactly: an unlearned model should be "
         "indistinguishable from a model trained on the dataset with the target "
         "records removed. For deep networks this exact criterion is "
         "unattainable in practice, and the field works with approximations "
         "whose quality must be measured empirically [18].")
    body(doc,
         "Three theoretical strands underpin this project. The first is "
         "knowledge localisation. Meng et al. [7] showed through causal tracing "
         "that specific factual associations in a transformer can be localised "
         "to identifiable layers and token positions, and that intervening there "
         "changes the model's factual output in a controlled way. MEMIT [22] "
         "extended this to large-scale editing. The same intervention machinery, "
         "run in reverse, is what allows erasure to be measured rather than "
         "inferred; this is the basis of the depth axis.")
    body(doc,
         "The second is the optimisation objective. Gradient ascent [5] simply "
         "maximises the loss on the data to be forgotten. It is effective but "
         "unstable, and if run too long it destroys the model. Gradient "
         "difference adds a retention term on data that must be preserved. "
         "Negative preference optimisation [4] recasts forgetting as the "
         "negative branch of a preference objective and was introduced "
         "specifically to avoid the catastrophic collapse that gradient ascent "
         "exhibits; SimNPO [23] simplifies it further.")
    body(doc,
         "The third is parameter arithmetic. Ilharco et al. [17] established "
         "that directions in parameter space behave compositionally: the "
         "difference between a fine-tuned and a base model forms a task vector "
         "that can be added or subtracted to add or remove a capability. UIPE "
         "[3] applies this to unlearning by extrapolating along the update "
         "direction. This project uses the same construction, but as an "
         "experimental dial rather than as a means of improving a score.")

    subheading(doc, "2.1.2  Existing Systems and Solutions")
    body(doc,
         "Benchmarks. TOFU [1] is the standard controlled setting for LLM "
         "unlearning and is the one adopted here, chiefly because it publishes "
         "retrained reference models. MUSE [9] evaluates on news and books "
         "across six criteria including sustainability under sequential "
         "requests. RWKU [10] targets real-world knowledge without a retain set. "
         "WMDP [8] measures hazardous capability and introduces a "
         "representation-level method, RMU. None of these benchmarks reports a "
         "mechanistic measure of erasure.")
    body(doc,
         "Methods. Beyond the gradient-based family, embedding-corrupted "
         "prompting [24] intervenes at the input rather than the weights; "
         "sparse-autoencoder subspace projection [26] intervenes in a learned "
         "feature basis; distillation has been shown to make unlearning "
         "substantially more robust [25]. These represent a shift from "
         "output-level to representation-level intervention, which increases the "
         "importance of representation-level measurement.")
    body(doc,
         "Evaluation and critique. A consistent line of work reports that "
         "unlearning is weaker than it appears. Lynch et al. [11] enumerate "
         "eight evaluation methods and show that methods passing some fail "
         "others. Łucki et al. [12] recover supposedly unlearned capability "
         "adversarially. Deeb and Roger [13] recover it by brief retraining. "
         "Doshi and Stickland [14] recover it through black-box probing alone. "
         "Hong et al. [15] and Lee et al. [16] examine the parametric side, the "
         "latter questioning whether localisation actually informs unlearning. "
         "Lee et al. [2] propose the Unlearning Depth Score, the first practical "
         "scalar measure of erasure depth, which this project implements and "
         "validates.")

    subheading(doc, "2.1.3  Research Findings for Existing Literature")
    body(doc,
         "Table 4 summarises the principal papers surveyed by the team. The "
         "survey was divided across members by theme; each member read and "
         "summarised the papers in their theme and contributed the corresponding "
         "rows.")
    table(doc, 4, "Literature survey summary",
          ["S. No.", "Name", "Paper Title", "Tools / Technology", "Findings",
           "Citation"],
          [
           ["1", "Team member 1", "TOFU: A Task of Fictitious Unlearning for LLMs",
            "TOFU dataset, Llama-2/3, Phi", "Controlled benchmark with retrained "
            "reference models; establishes the forget/retain split design used here",
            "[1]"],
           ["2", "Team member 1", "MUSE: Machine Unlearning Six-Way Evaluation",
            "News and books corpora", "Six evaluation criteria; shows methods "
            "degrade under sequential unlearning requests", "[9]"],
           ["3", "Team member 1", "RWKU: Benchmarking Real-World Knowledge Unlearning",
            "Real-entity probes", "Evaluates without a retain set; highlights "
            "difficulty of defining scope on real knowledge", "[10]"],
           ["4", "Team member 2", "The WMDP Benchmark",
            "RMU, representation control", "Hazardous-capability setting; "
            "introduces representation-level unlearning", "[8]"],
           ["5", "Team member 2", "Knowledge Unlearning for Mitigating Privacy Risks",
            "Gradient ascent", "Establishes gradient ascent as a baseline and "
            "documents its instability", "[5]"],
           ["6", "Team member 2", "Who's Harry Potter? Approximate Unlearning in LLMs",
            "Targeted fine-tuning", "Early demonstration of large-scale concept "
            "removal; evaluation is behavioural only", "[6]"],
           ["7", "Team member 3", "Negative Preference Optimization",
            "NPO objective, DPO branch", "Reframes forgetting as preference "
            "optimisation; explicitly addresses catastrophic collapse", "[4]"],
           ["8", "Team member 3", "Simplicity Prevails: Rethinking NPO",
            "SimNPO", "Simplifies NPO while retaining stability", "[23]"],
           ["9", "Team member 3", "LLM Unlearning via Embedding-Corrupted Prompts",
            "Input-space intervention", "Achieves forgetting without weight "
            "modification, separating behaviour from representation", "[24]"],
           ["10", "Team member 4", "Locating and Editing Factual Associations in GPT",
            "ROME, causal tracing", "Factual associations localise to specific "
            "layers and positions; source of the patching methodology", "[7]"],
           ["11", "Team member 4", "Mass-Editing Memory in a Transformer",
            "MEMIT", "Scales localised editing to thousands of facts", "[22]"],
           ["12", "Team member 4", "Editing Models with Task Arithmetic",
            "Task vectors", "Parameter-space differences compose; theoretical "
            "basis for the extrapolation dial", "[17]"],
           ["13", "Team member 5", "Measuring the Depth of LLM Unlearning",
            "Two-stage activation patching, UDS", "Defines the depth score "
            "implemented and validated in this project", "[2]"],
           ["14", "Team member 5", "UIPE: Unlearning with Parameter Extrapolation",
            "Parameter extrapolation", "Extrapolates the update vector; supplies "
            "the checkpoint-selection procedure adopted here", "[3]"],
           ["15", "Team member 5", "Eight Methods to Evaluate Robust Unlearning",
            "Multiple probes", "Methods that pass one evaluation fail others; "
            "motivates multi-axis measurement", "[11]"],
           ["16", "Team member 5", "An Adversarial Perspective on Machine Unlearning",
            "Adversarial recovery", "Supposedly unlearned capability is "
            "recoverable; evaluations overstate removal", "[12]"],
           ["17", "Team member 5", "Do Unlearning Methods Remove Information "
                                   "from Weights?", "Retraining audit",
            "Knowledge often remains in weights after apparent unlearning", "[13]"],
           ["18", "Team member 5", "Intrinsic Test of Unlearning Using "
                                   "Parametric Knowledge Traces",
            "ConceptVectors", "Parametric traces persist after behavioural "
            "unlearning", "[15]"],
          ],
          widths=[0.45, 0.85, 1.55, 1.0, 1.6, 0.55], font=9)

    subheading(doc, "2.1.4  Problem Identified")
    body(doc,
         "The survey converges on one conclusion. The field has many unlearning "
         "methods and several benchmarks, and a consistent set of critical "
         "results showing that the benchmarks overstate what the methods "
         "achieve. The reason is structural: the benchmarks measure behaviour, "
         "and behaviour can be changed without changing what the model knows. A "
         "mechanistic measure of depth now exists [2], and complementary "
         "parametric evidence exists [15], but neither has been related to how "
         "broadly forgetting generalises. There is therefore no way at present "
         "to state whether a given unlearning result is deep and narrow, broad "
         "and shallow, or genuinely both.")

    subheading(doc, "2.1.5  Survey of Tools and Technologies Used")
    body(doc,
         "PyTorch was selected as the modelling framework because activation "
         "patching requires forward hooks and direct manipulation of "
         "intermediate tensors, which it exposes cleanly. The Hugging Face "
         "Transformers library provides the model implementations and the "
         "TOFU-tuned checkpoints; Datasets provides the benchmark itself. SciPy "
         "supplies the rank-correlation statistics and Matplotlib the figures. "
         "Pytest is used for the test suite, supplemented by mutation testing, "
         "in which a deliberate fault is introduced into the implementation to "
         "confirm that some test detects it; a test that passes against a broken "
         "implementation provides no assurance. Section 3.4 records versions.")

    subheading(doc, "How this work builds on and differs from prior work")
    body(doc,
         "This project builds directly on three lines of work. The experimental "
         "setting, including the forget and retain splits and the retrained "
         "reference models, is taken unchanged from TOFU [1]. The depth "
         "measurement implements the Unlearning Depth Score of Lee et al. [2], "
         "including its layer-selection threshold and its weighted aggregation, "
         "and is validated by reproducing that paper's published values. The "
         "traversal mechanism implements the parameter extrapolation of UIPE "
         "[3], including the checkpoint-selection step of its algorithm, which "
         "in turn rests on task arithmetic [17].")
    body(doc,
         "It differs in three respects. First, prior work measures depth or "
         "behaviour, not both; this project measures both on the same models "
         "under the same conditions and on a common calibrated scale, which is "
         "what makes the depth-breadth relationship observable at all. Second, "
         "UIPE treats the extrapolation coefficient as a hyperparameter to be "
         "tuned for utility; here it is a controlled independent variable that "
         "generates a continuous trajectory through the plane. Third, breadth is "
         "reported against an empirically measured range bounded by the original "
         "and retrained models, rather than the nominal zero-to-one range used "
         "in the benchmarks surveyed, and the analysis refuses to draw a "
         "conclusion when that range is too small to support one.")

    heading(doc, "2.2  Software Requirement Specification")

    subheading(doc, "2.2.1  Introduction")
    para(doc, "2.2.1.1  Purpose", bold=True, space_before=6, space_after=2)
    body(doc,
         "This specification defines the requirements for the DeepErase "
         "measurement system: software that quantifies how deeply and how "
         "broadly knowledge has been removed from a language model by an "
         "unlearning procedure, and that traces the relationship between the two "
         "as unlearning strength is varied.")
    para(doc, "2.2.1.2  Intended Audience and Reading Suggestions", bold=True,
         space_before=6, space_after=2)
    body(doc,
         "The document is written for the project mentor and evaluation panel, "
         "for researchers who may reuse the measurement code, and for the "
         "project team. Readers seeking the scientific contribution should read "
         "Chapter 1 and Section 4.4. Readers concerned with implementation "
         "should read Chapter 3 and Sections 4.1 to 4.3.")
    para(doc, "2.2.1.3  Project Scope", bold=True, space_before=6, space_after=2)
    body(doc,
         "The system is a research instrument, not an end-user product. It is "
         "operated from the command line, consumes public benchmark data and "
         "public model checkpoints, and produces measurement artefacts. It does "
         "not perform unlearning as a service and makes no claim to remove "
         "knowledge on behalf of a third party.")

    subheading(doc, "2.2.2  Overall Description")
    para(doc, "2.2.2.1  Product Perspective", bold=True, space_before=6,
         space_after=2)
    body(doc,
         "The system is a self-contained Python package that sits on top of "
         "PyTorch and Hugging Face Transformers. It has no server component, no "
         "database and no persistent state beyond the run directories it writes. "
         "It is structured as a library with a thin command-line layer, so that "
         "each measurement can be invoked independently or composed into the "
         "full study. The architecture is described in Section 4.1 and shown in "
         "Figure 2.")
    para(doc, "2.2.2.2  Product Features", bold=True, space_before=6,
         space_after=2)
    for s in [
        "F1  Depth measurement. Computes the Unlearning Depth Score for a "
        "candidate model against a target and an oracle, including "
        "knowledge-encoding layer selection and weighted aggregation.",
        "F2  Breadth measurement. Scores forced-choice probes across tiers and "
        "calibrates the result against reference models.",
        "F3  Unlearning. Applies gradient ascent, gradient difference or "
        "negative preference optimisation, evaluating utility after each epoch "
        "and returning the best checkpoint above a utility floor.",
        "F4  Extrapolation. Constructs the update vector and produces models at "
        "any requested extrapolation strength, excluding non-trainable buffers.",
        "F5  Trajectory analysis. Assembles measurements into a trajectory, "
        "applies the utility control and the dynamic-range gate, and computes "
        "rank correlation.",
        "F6  Validation. Reproduces published reference values for the depth "
        "axis and compares against them within a stated tolerance.",
        "F7  Preflight. Verifies the installation and reports whether the "
        "planned run fits in available memory before it begins.",
        "F8  Reproducible output. Every run writes its configuration, raw "
        "measurements, derived results and figures to a timestamped directory.",
    ]:
        bullet(doc, s)

    subheading(doc, "2.2.3  External Interface Requirements")
    para(doc, "2.2.3.1  User Interfaces", bold=True, space_before=6,
         space_after=2)
    body(doc,
         "Interaction is entirely through a command-line interface. Five "
         "entry points are provided, corresponding to the use cases in "
         "Figure 8: preflight verification, depth validation, breadth "
         "measurement, annotation comparison, and the full depth-breadth study. "
         "Each accepts arguments for model scale, sample sizes and method "
         "selection, and each reports progress to the terminal while writing "
         "machine-readable artefacts to disk. There is no graphical interface; "
         "results are consumed as JSON and as generated figures.")
    para(doc, "2.2.3.2  Hardware Interfaces", bold=True, space_before=6,
         space_after=2)
    body(doc,
         "The system requires a CUDA-capable accelerator for the model-execution "
         "paths. Measurement alone requires memory for the resident model "
         "weights; at the 1B scale in bfloat16 this is approximately 2.5 GB per "
         "model. Full-parameter unlearning additionally requires memory for "
         "gradients and two optimiser moments, approximately four times the "
         "weight memory, so a 1B training run needs on the order of 11 GB "
         "including activations. The package computes this requirement before "
         "starting and refuses to begin a run that cannot complete. The test "
         "suite and all analysis code run on CPU without an accelerator.")
    para(doc, "2.2.3.3  Software Interfaces", bold=True, space_before=6,
         space_after=2)
    body(doc,
         "Python 3.10 or later; PyTorch with CUDA support; Hugging Face "
         "Transformers and Datasets for model and benchmark access; SciPy and "
         "NumPy for statistics; Matplotlib for figures; Pytest for the test "
         "suite. Model checkpoints and benchmark data are retrieved from the "
         "Hugging Face Hub and cached locally.")

    subheading(doc, "2.2.4  Other Non-functional Requirements")
    para(doc, "2.2.4.1  Performance Requirements", bold=True, space_before=6,
         space_after=2)
    body(doc,
         "A single depth measurement over one hundred examples across sixteen "
         "layers completes in approximately thirty seconds at the 1B scale. A "
         "breadth measurement over twelve hundred probe items completes in "
         "approximately ninety seconds. A complete study comprising unlearning "
         "and an eleven-point extrapolation sweep completes in under an hour at "
         "the 1B scale. The full automated test suite completes in under forty "
         "seconds on CPU, which is a deliberate design requirement: a suite that "
         "is slow to run is a suite that stops being run.")
    para(doc, "2.2.4.2  Safety Requirements", bold=True, space_before=6,
         space_after=2)
    body(doc,
         "The principal safety requirement is scientific rather than physical: "
         "the system must not report a measurement it cannot support. Three "
         "guards enforce this. Unlearning raises an error rather than returning "
         "a model if no checkpoint stayed above the utility floor, because a "
         "collapsed model scores maximally on every forget metric and would "
         "silently produce a meaningless result. Trajectory analysis reports "
         "retained utility alongside every point, so that general degradation "
         "cannot be mistaken for targeted forgetting. Analysis declines to state "
         "a finding when the measured range on either axis is too small to "
         "support one.")
    para(doc, "2.2.4.3  Security Requirements", bold=True, space_before=6,
         space_after=2)
    body(doc,
         "The system processes no personal data: all experimental subjects are "
         "fictitious individuals created for benchmarking. It requires no "
         "credentials, opens no network services and executes no user-supplied "
         "code. Downloaded model and dataset artefacts are obtained from a "
         "single well-known source and cached in a user-controlled directory.")

    heading(doc, "2.3  Cost Analysis")
    body(doc,
         "The project uses no licensed software and no paid data. The principal "
         "cost is computation, and Table 5 states it as a requirement in "
         "accelerator-hours so that it can be planned for irrespective of how "
         "the capacity is eventually supplied.")
    table(doc, 5, "Cost analysis",
          ["Item", "Basis", "Cost"],
          [["Benchmark data (TOFU)", "Publicly released for research use", "Nil"],
           ["Model checkpoints", "Publicly released under the Llama 3.2 "
                                 "community licence", "Nil"],
           ["Software stack", "PyTorch, Transformers, SciPy, Matplotlib, Pytest "
                              "— all open source", "Nil"],
           ["Development and testing", "Test suite and analysis run on CPU on "
                                       "team hardware", "Nil"],
           ["Compute requirement, work completed",
            "Depth validation, breadth calibration and one full study: "
            "approximately 25 accelerator-hours at the 1B scale",
            "25 GPU-h"],
           ["Compute requirement, remaining",
            "Replication across three methods and three seeds, plus the 3B "
            "scale: approximately 120 accelerator-hours",
            "120 GPU-h"],
           ["Storage requirement", "Model cache, run artefacts and checkpoints",
            "≈ 250 GB"],
           ["Total additional expenditure", "No licence or data purchase is "
                                            "required at any stage", "Nil"]],
          widths=[1.6, 3.0, 1.1])

    heading(doc, "2.4  Risk Analysis")
    body(doc,
         "Table 6 lists the identified risks with their assessed impact and "
         "likelihood. Risk R1 is the one that has most shaped the design: it is "
         "both high impact and high likelihood, because the failure it describes "
         "is silent. A model damaged by over-aggressive unlearning scores "
         "perfectly on every forgetting metric, so the failure does not announce "
         "itself and would be reported as an unusually good result. Its "
         "mitigation is accordingly built into the system as a hard guard rather "
         "than a monitoring practice, and Section 4.4.3 shows it operating.")
    table(doc, 6, "Risk analysis",
          ["No.", "Risk", "Impact", "Likelihood", "Mitigation"],
          [["R1", "Gradient ascent destroys the model instead of unlearning it, "
                  "and the damaged model scores maximally on every forget "
                  "metric.", "High", "High",
            "Utility is evaluated after every epoch and the returned checkpoint "
            "must remain above a utility floor; if none does, the run raises an "
            "error rather than returning a model."],
           ["R2", "Measurement resolution too coarse to detect the effect being "
                  "studied.", "High", "Medium",
            "Probe count chosen so that the smallest observable change is small "
            "relative to the expected effect; analysis is gated on measured "
            "dynamic range."],
           ["R3", "General degradation mistaken for targeted forgetting.",
            "High", "Medium",
            "A retained-knowledge control tier is measured at every point and "
            "reported alongside both axes."],
           ["R4", "Insufficient accelerator capacity for the remaining "
                  "replication.", "Medium", "Medium",
            "Memory requirement computed before each run; a parameter-efficient "
            "adaptation path reduces the training requirement to approximately "
            "weight-only memory if needed."],
           ["R5", "Implementation diverges from the published method it claims "
                  "to reproduce.", "High", "Low",
            "Depth axis validated against published reference values; a "
            "conformance document records the comparison equation by equation."],
           ["R6", "Tests pass without exercising the behaviour they name.",
            "Medium", "Medium",
            "Mutation testing: faults are deliberately introduced and the suite "
            "must detect them."],
           ["R7", "Scope expansion beyond the semester.", "Medium", "Medium",
            "Breadth tiers beyond paraphrase level and scales beyond 3B are "
            "explicitly out of scope for this evaluation."]],
          widths=[0.4, 1.5, 0.6, 0.7, 2.6])


# ===========================================================================
# CHAPTER 3
# ===========================================================================
def _chapter3(doc):
    chapter(doc, "Chapter 3  Methodology Adopted")

    heading(doc, "3.1  Investigative Techniques")
    body(doc,
         "Three investigative techniques were considered, summarised in "
         "Table 7. The choice is not cosmetic: it determines what kind of claim "
         "the project is entitled to make at the end.")
    table(doc, 7, "Investigative techniques considered",
          ["No.", "Technique", "Description", "Applicability to this project"],
          [["1", "Descriptive",
            "Scientific questions are investigated and observations of a "
            "phenomenon are recorded and catalogued.",
            "Rejected as the primary technique. The project does record "
            "observations, but a catalogue of measurements would not answer "
            "whether depth and breadth trade off; that requires a manipulated "
            "variable."],
           ["2", "Comparative",
            "Observations are made that compare two objects or phenomena.",
            "Rejected as the primary technique. Comparing unlearning methods "
            "against each other would confound the method with the strength at "
            "which it was applied, since different methods reach different "
            "operating points."],
           ["3", "Experimental",
            "An organised investigation including a control group, designed to "
            "test a hypothesis, with independent and dependent variables.",
            "Selected. The project has a manipulable independent variable, two "
            "dependent variables, and a control that distinguishes the effect of "
            "interest from a confound."]],
          widths=[0.4, 0.9, 1.9, 2.6])

    subheading(doc, "Justification of the experimental technique")
    body(doc,
         "The research question is causal in form: as an unlearning "
         "intervention is made stronger, what happens jointly to the depth and "
         "the breadth of forgetting? Answering it requires the strength of the "
         "intervention to be varied deliberately while everything else is held "
         "fixed, and requires an independent check that any observed change is "
         "attributable to forgetting rather than to something else. That is the "
         "definition of an experiment, and neither of the alternative techniques "
         "supplies it.")
    body(doc,
         "The independent variable is the extrapolation coefficient. After a "
         "model has been unlearned, the difference between its parameters and "
         "the original parameters defines a direction in parameter space. "
         "Scaling that direction by a coefficient produces a family of models "
         "that differ only in how far they have been pushed along one single "
         "path. This matters, because the alternative way of varying unlearning "
         "strength, namely training for different numbers of epochs, changes the "
         "direction of the update as well as its magnitude, and therefore "
         "confounds strength with trajectory. Extrapolation holds the direction "
         "fixed and varies only the distance, which is exactly the control an "
         "experiment requires.")
    body(doc,
         "The dependent variables are depth and breadth. Both are measured "
         "independently at every point on the path. They are constructed to be "
         "comparable: each is anchored at 0 for the original model and 1 for a "
         "model retrained without the target data, so that a value of 0.5 means "
         "the same thing on both axes.")
    body(doc,
         "The control is retained-knowledge utility, measured on facts that "
         "must not be forgotten. This control does the work that makes the "
         "experiment interpretable. Any intervention strong enough to damage a "
         "model will reduce its score on the forget set, and a sufficiently "
         "damaged model scores perfectly on every forgetting metric while having "
         "learned nothing about forgetting. Without the control, general damage "
         "and targeted removal are indistinguishable. With it, a point at which "
         "utility has fallen can be identified as such and excluded or reported "
         "as a limitation.")
    body(doc,
         "Two further controls follow from the experimental design. The first is "
         "the pair of reference models, which fix the ends of both scales "
         "empirically rather than by assumption; without them a raw score has no "
         "interpretable zero. The second is validation of the instrument itself "
         "against published values before it is used to generate new results, "
         "reported in Section 4.4. An instrument that has not been checked "
         "against a known quantity cannot support a claim about an unknown one.")
    body(doc,
         "Finally, the experimental frame determines how the results may be "
         "read, and imposes a restriction that the project accepts explicitly. "
         "The points along one extrapolation path are successive evaluations of "
         "a single continuous trajectory, not independent samples drawn from a "
         "population. A significance test computed over them would be invalid, "
         "because its assumption of independence is violated by construction and "
         "because evaluating the path more densely would appear to strengthen "
         "the evidence without adding any. Rank correlation is therefore "
         "reported as a description of the shape of the trajectory, and any "
         "claim about the depth-breadth relationship is deferred until it has "
         "been replicated across independent runs with different methods and "
         "different random seeds.")

    heading(doc, "3.2  Proposed Solution")
    body(doc,
         "The solution is a measurement pipeline in five stages, implemented as "
         "the module structure shown in Figure 2 and executed in the order shown "
         "in Figure 6.")

    subheading(doc, "Stage 1: Anchoring")
    body(doc,
         "Two reference models are loaded and measured: the model fine-tuned on "
         "all two hundred authors, and the model retrained with the forget "
         "authors excluded. The first knows the target facts with certainty; the "
         "second has certainly never seen them. Measuring both on the breadth "
         "probes yields the ceiling and floor of the achievable range. This step "
         "is what allows a later score to be expressed as a fraction of the "
         "distance to genuine removal rather than as a raw number whose "
         "interpretation depends on the difficulty of the probe set.")

    subheading(doc, "Stage 2: Depth calibration")
    body(doc,
         "Depth is measured by two-stage activation patching, illustrated in "
         "Figure 3. In the first stage, hidden states are captured from the "
         "retrained model and inserted into the model that knows the fact, one "
         "layer at a time, at the token positions immediately preceding the "
         "target entity. The resulting drop in the log-probability the model "
         "assigns to that entity, written ΔS₁ for each layer, measures how much "
         "of the fact that layer carries. Layers whose ΔS₁ does not exceed a "
         "threshold of τ = 0.05 nats are discarded as carrying no measurable "
         "part of the fact; the remainder are the knowledge-encoding layers, and "
         "only these are used subsequently. Discarding them is not an "
         "optimisation but a necessity: the depth score divides by ΔS₁, and a "
         "layer that carries nothing would contribute a ratio of two "
         "indistinguishable noise terms. This stage depends only on the two "
         "reference models, so it is computed once and cached.")

    subheading(doc, "Stage 3: Unlearning with checkpoint selection")
    body(doc,
         "An unlearning method is applied to the model that knows the facts. "
         "Three are implemented: gradient ascent, which maximises loss on the "
         "forget set; gradient difference, which adds a retention term; and "
         "negative preference optimisation, which treats the forget set as the "
         "negative branch of a preference objective against a frozen reference "
         "copy.")
    body(doc,
         "Training does not simply run to a fixed number of epochs. Following "
         "the checkpoint-selection step of the UIPE algorithm [3], utility on "
         "retained knowledge and loss on the forget set are both evaluated after "
         "every epoch, and the model returned is the checkpoint that forgot the "
         "most while retaining at least ninety per cent of its starting utility. "
         "If no checkpoint satisfies that condition the procedure raises an "
         "error rather than returning a model. This is not a refinement but a "
         "requirement: gradient ascent left to run freely destroys the model, "
         "and a destroyed model scores maximally on every forgetting metric, so "
         "the failure is invisible to every downstream measurement. The "
         "behaviour of this stage on real data is shown in Figure 10, and the "
         "state machine it implements is Figure 7.")

    subheading(doc, "Stage 4: Traversal by extrapolation")
    body(doc,
         "The update vector is the elementwise difference between the unlearned "
         "and original parameters. A model at extrapolation strength α is "
         "constructed by adding α times that vector to the unlearned parameters, "
         "so α = 0 recovers the unlearned model and α = 1 applies the update a "
         "second time in full. Non-trainable buffers are excluded by name: these "
         "hold quantities such as cached rotary embeddings, which are not "
         "learned parameters and whose extrapolation would corrupt the model "
         "while appearing to be a legitimate weight update.")

    subheading(doc, "Stage 5: Measurement and analysis")
    body(doc,
         "Every model on the path is scored on three quantities. Depth repeats "
         "the patching procedure of Stage 2 with the extrapolated model as the "
         "source, taking the ratio of its per-layer effect to the reference "
         "effect, clipping to the unit interval, and aggregating across "
         "knowledge-encoding layers weighted by how much each layer carries. "
         "Breadth scores the forced-choice probes and converts the raw leakage "
         "into the calibrated range established in Stage 1. Utility scores the "
         "retain control tier. The three are assembled into a trajectory, and "
         "the analysis applies the guards described in Section 2.2.4.2 before "
         "reporting anything.")

    heading(doc, "3.3  Work Breakdown Structure")
    body(doc,
         "The work is decomposed into twelve packages, listed in Table 8 and "
         "scheduled in Figure 1. Each package produces something that can be "
         "run and tested on its own, so that progress is demonstrable rather "
         "than merely reported.")
    figure(doc, "fig08_wbs", 1,
           "Work breakdown structure and schedule. Completed packages are shown "
           "solid; planned packages are shown in grey, with the mid-semester "
           "evaluation marked.", width=6.0)
    body(doc,
         "Figure 1 shows that the eight packages carrying the measurement "
         "system, W1 to W9, are complete as of the mid-semester evaluation "
         "marked on the schedule, and that the remaining effort is replication "
         "and extension rather than new construction. The ordering is not "
         "arbitrary: the depth axis was validated (W5) before it was used to "
         "generate any new result, and the breadth axis was calibrated (W6) "
         "before either axis was applied to an unlearned model, so that no "
         "measurement in Chapter 4 rests on an instrument that had not first "
         "been checked against a known quantity.")
    table(doc, 8, "Work breakdown structure",
          ["No.", "Work package", "Deliverable", "Status"],
          [["W1", "Literature review and domain study",
            "Survey of 60 papers; identified research gaps", "Complete"],
           ["W2", "Environment, package skeleton, test harness",
            "Installable package; continuous test suite", "Complete"],
           ["W3", "TOFU data layer and entity spans",
            "Split loading, balanced sampling, span location", "Complete"],
           ["W4", "Activation patching and depth backend",
            "Hidden-state capture, patched forward, UDS computation", "Complete"],
           ["W5", "Depth-axis validation",
            "Reproduction of published reference values", "Complete"],
           ["W6", "Breadth axis and calibration",
            "Tiered forced-choice scoring; reference calibration", "Complete"],
           ["W7", "Unlearning methods and checkpoint selection",
            "GA, GradDiff, NPO with utility-floor stopping", "Complete"],
           ["W8", "Extrapolation dial and trajectory analysis",
            "Update vector, α sweep, trajectory guards", "Complete"],
           ["W9", "First end-to-end trajectory",
            "Complete depth-breadth path with controls", "Complete"],
           ["W10", "Replication across methods and seeds",
            "NPO and GradDiff arms; three seeds each", "In progress"],
           ["W11", "Scale-up and additional breadth tiers",
            "3B scale; alias and entailment tiers", "Planned"],
           ["W12", "Analysis, write-up and submission",
            "Final report and artefacts", "Planned"]],
          widths=[0.45, 1.9, 2.4, 0.85])

    heading(doc, "3.4  Tools and Technology")
    body(doc,
         "Table 9 lists the tools used and the role each plays. Two selections "
         "were made on technical grounds rather than convenience. PyTorch was "
         "required because activation patching depends on registering forward "
         "hooks and substituting intermediate tensors mid-computation, which a "
         "framework that treats the forward pass as opaque would not permit. "
         "Pytest was supplemented by mutation testing because the correctness "
         "properties at stake here are numerical and easy to satisfy vacuously; "
         "a test that passes against a deliberately broken implementation "
         "provides no assurance, and several early tests were found to do "
         "exactly that.")
    table(doc, 9, "Tools and technology",
          ["Category", "Tool", "Version", "Role in the project"],
          [["Language", "Python", "3.10+",
            "Implementation language for the entire package"],
           ["Deep learning", "PyTorch", "2.x",
            "Model execution; forward hooks for activation capture and patching"],
           ["Models", "Hugging Face Transformers", "4.x",
            "Llama-3.2 implementations and TOFU-tuned checkpoints"],
           ["Data", "Hugging Face Datasets", "3.x",
            "TOFU benchmark splits and perturbed answer sets"],
           ["Numerics", "NumPy", "1.26+", "Array operations in the analysis layer"],
           ["Statistics", "SciPy", "1.13+",
            "Spearman rank correlation for trajectory description"],
           ["Visualisation", "Matplotlib", "3.9+",
            "All figures in this report and in run artefacts"],
           ["Testing", "Pytest", "8.x",
            "489-test automated suite, verified by mutation testing"],
           ["Environment", "Conda", "24.x",
            "Reproducible dependency isolation"],
           ["Version control", "Git", "2.x",
            "Source history and result traceability"]],
          widths=[1.0, 1.5, 0.7, 2.5])


# ===========================================================================
# CHAPTER 4
# ===========================================================================
def _chapter4(doc):
    chapter(doc, "Chapter 4  Design Specifications")

    heading(doc, "4.1  System Architecture")
    body(doc,
         "The system is organised as a layered pipeline, shown in Figure 2. "
         "Data flows downward: benchmark data and model checkpoints enter at the "
         "top, are transformed by the preparation, unlearning and extrapolation "
         "modules, are scored independently by the two measurement axes, and are "
         "combined in the analysis layer into artefacts written to disk.")
    figure(doc, "fig01_architecture", 2,
           "System architecture of the DeepErase measurement pipeline. Data "
           "flows from the input layer through preparation, unlearning and "
           "extrapolation into two independent measurement axes, which are "
           "combined by the analysis layer.")
    body(doc,
         "Figure 2 shows four properties of the design that were deliberate "
         "rather than incidental. First, the two measurement axes are entirely "
         "independent of one another: the depth module has no knowledge of the "
         "breadth module and neither shares state with the other, so a fault in "
         "one cannot silently propagate into the other's result. Second, the "
         "unlearning and extrapolation modules are separated, which is what "
         "permits a single unlearned model to generate an entire trajectory "
         "without retraining. Third, the input layer is the only component that "
         "touches the network, so every run below it is deterministic given its "
         "cached inputs. Fourth, the analysis layer is the sole writer of "
         "conclusions, which concentrates the scientific guards described in "
         "Section 2.2.4.2 in one reviewable place.")
    body(doc,
         "The depth path, shown on the left of Figure 2, consumes prepared "
         "examples with located entity spans and produces a scalar per candidate "
         "model. The breadth path on the right consumes probe items grouped into "
         "tiers and produces a raw leakage figure that the calibration converts "
         "into the common scale. Both paths depend on the reference models "
         "measured once at the start of a run.")

    heading(doc, "4.2  Design Level Diagrams")

    subheading(doc, "4.2.1  Depth measurement procedure")
    body(doc,
         "Figure 3 details the two-stage patching procedure that constitutes "
         "the depth measurement, expanding the left-hand block of Figure 2.")
    figure(doc, "fig02_patching", 3,
           "Two-stage activation patching used to compute the depth score. "
           "Stage 1 calibrates how much of the target fact each layer carries; "
           "Stage 2 measures how much of that the candidate model has lost.")
    body(doc,
         "The essential idea in Figure 3 is that the same intervention is "
         "performed twice with different sources. In Stage 1 the source is the "
         "retrained oracle, which is known not to hold the fact; the resulting "
         "drop ΔS₁ therefore measures the full effect of removing that layer's "
         "contribution, and serves as the denominator. In Stage 2 the source is "
         "the candidate model under test; the drop ΔS₂ measures how much of that "
         "same contribution the candidate has lost. Their ratio, clipped to the "
         "unit interval, is the fraction of the layer's knowledge that "
         "unlearning removed. Aggregating across knowledge-encoding layers, "
         "weighted by ΔS₁ so that layers carrying more of the fact count for "
         "more, gives the score. The construction guarantees the two fixed "
         "points quoted throughout this report: a model identical to the target "
         "scores 0, and a model identical to the oracle scores 1.")

    subheading(doc, "4.2.2  Breadth probe structure")
    body(doc,
         "Figure 4 shows how a single fact is probed at increasing distance "
         "from the form in which the model was trained.")
    figure(doc, "fig03_breadth_tiers", 4,
           "Breadth tier structure: one fact probed at increasing semantic "
           "distance, with a retain control tier that must not change.",
           width=5.8)
    body(doc,
         "Tier B0 in Figure 4 uses the exact question form the model saw during "
         "fine-tuning, and therefore measures forgetting in the narrowest "
         "possible sense. Tier B1 asks for the same fact in different words, so "
         "a model that scores well on B0 but poorly on B1 has suppressed a "
         "surface form rather than removed a fact. Tier R is the control: it "
         "asks about authors that were never in the forget set, and its score "
         "must remain unchanged. A method that appears to succeed on B0 and B1 "
         "while R also falls has damaged the model generally rather than "
         "forgotten selectively, and Figure 4 makes the necessity of that "
         "control explicit. Tiers covering aliases, entailed facts and multi-hop "
         "inference are designed but not yet populated, and are listed in "
         "Section 5.4.")

    subheading(doc, "4.2.3  Class diagram")
    body(doc,
         "Figure 5 shows the principal classes and their relationships.")
    figure(doc, "fig05_class", 5,
           "Class diagram of the principal modules, showing the configuration, "
           "measurement and analysis layers.")
    body(doc,
         "The upper row of Figure 5 is the configuration and resource layer. "
         "RunConfig captures everything that defines a run, so that a saved "
         "configuration is sufficient to repeat it. ModelManager owns model "
         "lifetime, which matters because several models are needed at different "
         "times and memory must be released between phases. MemoryPlan computes "
         "whether a planned run fits before it begins. The middle row holds the "
         "three measurement components. The lower row is the analysis layer: "
         "unlearn produces an UnlearnHistory recording per-epoch evaluations; "
         "each scored extrapolation point becomes a PlanePoint; and a Trajectory "
         "aggregates many PlanePoints and applies the guards before any "
         "conclusion is drawn.")

    subheading(doc, "4.2.4  Sequence diagram")
    body(doc,
         "Figure 6 traces the order of operations in a complete study run, the "
         "most significant use case in the system.")
    figure(doc, "fig06_sequence", 6,
           "Sequence diagram for a complete depth-breadth study run, from "
           "reference measurement through to written artefacts.")
    body(doc,
         "Two aspects of Figure 6 are worth drawing out. The reference models "
         "are measured and then explicitly freed before unlearning begins, "
         "because the training phase requires several times the memory of the "
         "measurement phase and the two must not overlap. The self-directed "
         "arrows on the unlearn lifeline are the per-epoch evaluation and "
         "checkpoint-selection loop of Stage 3; they occur inside a single call "
         "and are the reason that call can fail rather than return a damaged "
         "model. The extrapolation loop that follows reuses one unlearned model "
         "for all eleven measured points, which is what makes the trajectory a "
         "controlled path rather than eleven unrelated training runs.")

    subheading(doc, "4.2.5  State chart")
    body(doc,
         "Figure 7 gives the state chart for a model as it passes through the "
         "system, which is the significant object in this project.")
    figure(doc, "fig07_statechart", 7,
           "State chart for the model lifecycle, showing per-epoch evaluation, "
           "checkpoint selection, the abort path, and the extrapolation loop.")
    body(doc,
         "The state chart in Figure 7 makes the safety property of Section "
         "2.2.4.2 explicit. From the Evaluated state there are exactly two "
         "transitions: a checkpoint whose utility has fallen below the floor "
         "moves to Rejected and can never be selected, and one that remains "
         "above the floor becomes an accepted candidate. If every epoch reaches "
         "Rejected, the machine terminates in the abort state rather than "
         "reaching Selected; there is no path by which a collapsed model becomes "
         "the returned result. The loop at the lower right is the extrapolation "
         "sweep, which revisits the measurement states once per value of α.")

    heading(doc, "4.3  User Interface Diagrams")
    body(doc,
         "The system is operated from the command line and has no graphical "
         "interface, so the interface is characterised by its use cases rather "
         "than by screen layouts. Figure 8 gives the use case diagram for the "
         "overall system.")
    figure(doc, "fig04_usecase", 8,
           "Use case diagram for the overall system, showing the researcher and "
           "evaluation actors and the seven supported operations.",
           width=5.2)
    body(doc,
         "The researcher actor in Figure 8 drives all seven use cases. The "
         "evaluation actor is a reader of results rather than an operator, and "
         "interacts only with the validation, measurement and export cases; this "
         "distinction is why the system writes machine-readable artefacts for "
         "every run rather than printing to the terminal alone. Each use case "
         "corresponds to one command-line entry point. The preflight case is "
         "separated deliberately: it verifies the installation and computes "
         "whether the requested run fits in available memory before any long "
         "operation begins, so that a configuration error is reported in seconds "
         "rather than after an hour of computation.")

    heading(doc, "4.4  Snapshots of Working Prototype")
    body(doc,
         "This section reports what the implemented system produces. The results "
         "are presented in the order in which the instruments were built and "
         "checked: first the validation of the depth axis against known values, "
         "then the calibration of the breadth axis, then the behaviour of the "
         "unlearning stage, and finally the complete trajectory that the "
         "assembled system generates.")

    subheading(doc, "4.4.1  Validation of the depth axis")
    body(doc,
         "Before the depth measurement could be used on new models it was "
         "checked against the published reference values of Lee et al. [2], "
         "which report the score for three retained models whose relationship to "
         "the forget set is known exactly. Table 10 and Figure 9 give the "
         "comparison.")
    table(doc, 10, "Depth-metric validation against published values",
          ["Model scored", "Relationship to the forget set", "Published [2]",
           "This implementation", "Deviation"],
          [["full", "Learned every target fact", "0.002", "0.000", "0.002"],
           ["retain99", "Never saw 10% of the forget set", "0.153", "0.126",
            "0.027"],
           ["retain95", "Never saw 50% of the forget set", "0.496", "0.486",
            "0.010"],
           ["retain90", "Never saw the entire forget set", "1.000", "1.000",
            "0.000"]],
          widths=[1.0, 1.9, 0.9, 1.0, 0.8])
    figure(doc, "fig09_uds_validation", 9,
           "Depth-metric validation against published reference values. The "
           "implementation reproduces the published ordering and magnitudes "
           "within the stated tolerance.", width=4.6)
    body(doc,
         "Table 10 and Figure 9 show agreement well inside the tolerance of "
         "0.08 threshold we set in advance: the mean absolute deviation is 0.010. The "
         "ordering is reproduced exactly, and both endpoints are exact by "
         "construction, since a model patched into itself produces no change and "
         "a model identical to the calibration source produces the full change. "
         "The result held across seven runs at two model scales and under "
         "variation in prompt format, example sampling and entity-span policy, "
         "which establishes that the measurement is not an artefact of any one "
         "configuration choice. On this basis the depth axis is treated as a "
         "validated instrument in the remainder of the report.")

    subheading(doc, "4.4.2  Calibration of the breadth axis")
    body(doc,
         "The breadth axis was calibrated by measuring the two reference models "
         "on the probe set. Table 11 gives the result.")
    table(doc, 11, "Breadth-axis calibration measurements",
          ["Model", "Knowledge of forget authors", "Raw leakage on forget tiers",
           "Score on retain control", "Calibrated breadth"],
          [["full", "Present", "0.771", "0.797", "0.000"],
           ["retain90", "Absent", "0.524", "0.807", "1.000"],
           ["", "Usable dynamic range", "0.248", "", ""]],
          widths=[0.8, 1.4, 1.3, 1.1, 1.0])
    body(doc,
         "Table 11 establishes two things. First, the measurement discriminates "
         "correctly: the model that learned the authors scores 0.771 on them, "
         "essentially its score on knowledge it certainly holds, while the model "
         "that never saw them scores 0.524, far below its own retain score of "
         "0.807. The instrument responds to the presence of knowledge and not to "
         "some general property of the questions. Second, the achievable range "
         "is 0.248 rather than the nominal 1.000. A model with no knowledge of a "
         "subject still rejects many alternatives because they are implausible "
         "on general grounds, so the floor sits well above chance. Reporting raw "
         "numbers as though 0 and 1 were the endpoints would understate every "
         "effect by roughly a factor of four, and this is the concrete instance "
         "of research gap G5. All breadth values in this report are calibrated "
         "to the measured range.")

    subheading(doc, "4.4.3  Behaviour of the unlearning stage")
    body(doc,
         "Figure 10 shows the unlearning stage operating on the forget set, "
         "with the checkpoint-selection rule of Stage 3 active.")
    figure(doc, "fig10_training", 10,
           "Measured unlearning trajectory. Forget-set loss rises steadily while "
           "utility is monitored against a floor; epochs beyond the floor are "
           "rejected and the checkpoint at epoch 12 is selected.", width=5.4)
    body(doc,
         "Figure 10 is the empirical justification for a design decision that "
         "might otherwise look like excess caution. Forget-set loss rises "
         "monotonically throughout, from 1.60 to 9.83, so by the usual "
         "behavioural criterion every epoch represents better forgetting than "
         "the one before. Utility, however, holds near its starting value for "
         "twelve epochs and then falls away sharply, reaching 0.45 by epoch 19 "
         "against a starting value of 0.80. The last four epochs describe a "
         "model that has been damaged, not one that has forgotten selectively, "
         "and every forgetting metric would score them as the best results of "
         "the run. The selection rule returns epoch 12, the final checkpoint "
         "above the floor. Without the control there would be no signal in the "
         "forget-set curve to indicate that anything had gone wrong.")

    subheading(doc, "4.4.4  Complete depth-breadth trajectory")
    body(doc,
         "The assembled system was run end to end: unlearning by gradient ascent "
         "with checkpoint selection, followed by an eleven-point extrapolation "
         "sweep with depth, breadth and utility measured at each point. Table 12 "
         "and Figure 11 give the result. The complete run is stored as "
         "study_ga_1B_20260816_165423, and every figure in this section is "
         "generated from the measurements in that directory.")
    table(doc, 12, "Measured depth-breadth trajectory",
          ["α", "Breadth", "Depth", "Retained utility"],
          [["0.0", "0.237", "0.353", "0.720"],
           ["0.1", "0.237", "0.424", "0.710"],
           ["0.2", "0.333", "0.480", "0.695"],
           ["0.3", "0.333", "0.531", "0.677"],
           ["0.4", "0.354", "0.572", "0.657"],
           ["0.5", "0.399", "0.608", "0.652"],
           ["0.6", "0.429", "0.639", "0.650"],
           ["0.7", "0.485", "0.666", "0.647"],
           ["0.8", "0.510", "0.690", "0.640"],
           ["0.9", "0.551", "0.708", "0.632"],
           ["1.0", "0.571", "0.725", "0.625"]],
          widths=[0.9, 1.3, 1.3, 1.6])
    figure(doc, "fig11_plane", 11,
           "Measured depth-breadth trajectory under extrapolation. Panel (a) "
           "shows the path through the plane; panel (b) shows each axis and the "
           "utility control against the extrapolation coefficient.", width=6.2)
    body(doc,
         "Panel (a) of Figure 11 shows the path traced through the plane as the "
         "extrapolation coefficient increases from 0 to 1. Panel (b) separates "
         "the three measured quantities. Both axes rise together: depth moves "
         "from 0.353 to 0.725 and breadth from 0.237 to 0.571, so at the "
         "strongest setting the model has achieved roughly seventy-three per "
         "cent of the depth and fifty-seven per cent of the breadth of a model "
         "retrained without the data. The path sits consistently above the "
         "diagonal in panel (a), meaning that at every point this method has "
         "removed the knowledge more deeply than it has removed it widely.")
    body(doc,
         "Three qualifications apply to this result, and all three are recorded "
         "here rather than in the conclusions because they bear on how the "
         "figure should be read.")
    body(doc,
         "First, the utility control does not hold constant across the sweep. "
         "Retained-knowledge utility declines from 0.720 to 0.625 as α "
         "increases, a fall of 0.095. Part of the joint rise in the two axes is "
         "therefore attributable to general degradation rather than to targeted "
         "forgetting, and the trajectory cannot be interpreted as a pure "
         "measurement of forgetting until the two contributions are separated. "
         "This is the single most important open item and is addressed first in "
         "Section 5.4.")
    body(doc,
         "Second, the points are not independent observations. They are eleven "
         "evaluations of one continuous path, so the rank correlation between "
         "the axes describes the shape of that path and nothing more; no "
         "significance can be attached to it, for the reasons given in "
         "Section 3.1.")
    body(doc,
         "Third, this is one method, one model, one random seed. The trajectory "
         "demonstrates that the assembled system produces a coherent, "
         "interpretable measurement across a substantial span of both axes, "
         "which is what this stage of the project set out to establish. It is "
         "not yet an answer to the research question, and Section 5.4 sets out "
         "the replication required before one is offered.")

    subheading(doc, "4.4.5  Verification of the implementation")
    body(doc,
         "Table 13 summarises the automated test suite. Coverage is reported by "
         "component rather than by line count, because the relevant question is "
         "whether each scientific claim the code makes is checked, not what "
         "fraction of statements execute.")
    table(doc, 13, "Automated test suite coverage",
          ["Test file", "Tests", "Representative properties verified"],
          [["test_config.py", "50",
            "Run requirements computed before execution; a run that cannot "
            "complete is refused rather than started"],
           ["test_tofu.py", "58",
            "Balanced sampling across the nested forget splits; prompt "
            "construction; entity-span location"],
           ["test_reference_spans.py", "23",
            "Published annotations located by character offset rather than by "
            "tokeniser-specific indices"],
           ["test_patching.py", "84",
            "Hidden-state capture and insertion at specified layers and token "
            "positions; no leakage between models"],
           ["test_metrics.py", "84",
            "Layer selection, ratio clipping, weighted aggregation; exact "
            "endpoints; trajectory guards"],
           ["test_breadth.py", "37",
            "Forced-choice ranking; calibration against reference models; "
            "scoring against the correct reference answer"],
           ["test_unlearn.py", "51",
            "Objective signs; checkpoint selection against the utility floor; "
            "abort on total collapse"],
           ["test_extrapolation.py", "42",
            "Update vector construction; buffer exclusion by name; scaling "
            "behaviour"],
           ["test_models.py", "39",
            "Device resolution; memory release between phases"],
           ["test_runner.py", "21",
            "Run directories written completely; interrupted runs resume"],
           ["Total", "489", "All passing; verified by mutation testing"]],
          widths=[1.7, 0.6, 3.4])
    body(doc,
         "The suite is verified by mutation testing: a fault is deliberately "
         "introduced into the implementation and the suite is re-run, and if it "
         "still passes then the corresponding test was not exercising the "
         "behaviour it claimed to. This check was applied to each scientific "
         "guard described in Section 2.2.4.2 and to the depth and breadth "
         "computations, and in each case the introduced fault was detected.")


# ===========================================================================
# CHAPTER 5
# ===========================================================================
def _chapter5(doc):
    chapter(doc, "Chapter 5  Conclusions and Future Scope")

    heading(doc, "5.1  Work Accomplished")
    body(doc,
         "Progress is stated against the objectives approved at the proposal "
         "evaluation and listed in Table 3.")
    subheading(doc, "O1  Reproducible experimental environment and data pipeline")
    body(doc,
         "Complete. The project is an installable Python package with pinned "
         "dependencies and a 489-test automated suite. The data layer loads the "
         "TOFU splits, samples them in a way that accounts for their nested "
         "structure, and locates target entities within answers either "
         "heuristically or from published annotations. Every run writes a "
         "self-describing directory containing its configuration, raw "
         "measurements and derived results.")
    subheading(doc, "O2  Mechanistic depth measurement, validated")
    body(doc,
         "Complete. Two-stage activation patching is implemented in full, "
         "including knowledge-encoding layer selection and weighted "
         "aggregation. It reproduces published reference values to a mean "
         "absolute deviation of 0.010, against our pre-set 0.08 threshold, across seven "
         "runs and two model scales (Table 10, Figure 9).")
    subheading(doc, "O3  Breadth measurement, calibrated")
    body(doc,
         "Complete for the exact, paraphrase and control tiers. The measurement "
         "discriminates correctly between models that hold and do not hold the "
         "target knowledge, and is calibrated against those models rather than "
         "reported on a nominal scale; the measured range is 0.248 (Table 11). "
         "The alias, entailment and multi-hop tiers are designed but not "
         "populated.")
    subheading(doc, "O4  Unlearning methods with a principled stopping rule")
    body(doc,
         "Complete. Gradient ascent, gradient difference and negative preference "
         "optimisation are implemented. Utility is evaluated after every epoch "
         "and the returned checkpoint must remain above a utility floor; if none "
         "does, the procedure aborts. Figure 10 shows why this is necessary "
         "rather than optional.")
    subheading(doc, "O5  Traversal of the depth-breadth plane")
    body(doc,
         "First trajectory obtained. The complete pipeline runs end to end and "
         "produces an eleven-point path with both axes and the utility control "
         "measured at every point (Table 12, Figure 11). Both axes span a "
         "substantial fraction of their calibrated range. Characterisation of "
         "the relationship awaits the replication in Section 5.4.")
    subheading(doc, "O6  Reporting across methods, seeds and scales")
    body(doc,
         "In progress. The gradient ascent arm at the 1B scale is complete. The "
         "negative preference optimisation and gradient difference arms, the "
         "additional seeds and the 3B scale remain.")

    heading(doc, "5.2  Conclusions")
    body(doc,
         "The work completed supports four conclusions.")
    body(doc,
         "The depth of unlearning can be measured reliably. Reproducing "
         "published reference values to a mean absolute deviation of 0.010, "
         "across two model scales and under variation in prompt format, "
         "sampling and annotation policy, establishes that the implementation "
         "measures the intended quantity rather than an artefact of one "
         "configuration.")
    body(doc,
         "Behavioural measures require calibration before they can be "
         "interpreted. A model that has never encountered a set of facts still "
         "scores 0.524 on probes about them, against a ceiling of 0.771 for a "
         "model that learned them. The usable range is 0.248, roughly a quarter "
         "of the nominal range, and results reported on the nominal scale "
         "understate effects accordingly.")
    body(doc,
         "A stopping rule is a requirement of unlearning, not a refinement. "
         "Forget-set loss rises monotonically for the entire duration of "
         "training, including throughout the period in which the model is being "
         "destroyed. Every purely behavioural forgetting metric scores the "
         "damaged model as the best result. Only a control on retained knowledge "
         "distinguishes the two, which means that any unlearning result reported "
         "without such a control is uninterpretable.")
    body(doc,
         "Depth and breadth are separately measurable on a common scale. This "
         "was the enabling result the project needed, and it has been "
         "demonstrated: the two axes are computed independently, are anchored to "
         "the same pair of reference models, and both traverse a substantial "
         "fraction of their range under a controlled intervention. In the single "
         "trajectory measured so far the two rise together, with depth "
         "consistently ahead of breadth; whether that reflects a property of "
         "unlearning or of gradient ascent specifically cannot be determined "
         "from one method, one seed and a run in which the utility control did "
         "not hold constant. That determination is the work of the remainder of "
         "the project.")

    heading(doc, "5.3  Environmental, Economic and Social Benefits")
    subheading(doc, "Environmental")
    body(doc,
         "Retraining a large model to remove data consumes energy on the scale "
         "of the original training run, and must be repeated for every removal "
         "request. Unlearning is orders of magnitude cheaper, but is only worth "
         "substituting if it works. By making the strength of an unlearning "
         "claim measurable, this work supports the substitution where it is "
         "justified and identifies where it is not, which is the condition for "
         "the energy saving to be real rather than nominal. The project's own "
         "footprint is modest: measurements use published checkpoints and the 1B "
         "and 3B scales, and the analysis and test suite run on CPU.")
    subheading(doc, "Economic")
    body(doc,
         "The cost of compliance with a data erasure request is currently "
         "either very high, if the model is retrained, or unquantified, if an "
         "unlearning method is used without adequate verification. Measurement "
         "that distinguishes removal from suppression allows organisations to "
         "adopt the cheaper route where the evidence supports it, and gives "
         "auditors a basis on which to accept or reject the claim. The "
         "instruments developed here are released as open source, so the cost of "
         "adopting them is limited to compute.")
    subheading(doc, "Social")
    body(doc,
         "A right to erasure that cannot be verified is not a right in any "
         "practical sense. When an organisation states that personal "
         "information has been removed from a deployed model, the person "
         "concerned currently has no way to assess that claim, and the evidence "
         "surveyed in Section 1.1 indicates that such claims are frequently "
         "stronger than the underlying method supports. Work that makes the "
         "distinction measurable serves the interest of the individuals whose "
         "data is at stake, and supports the enforcement of data protection "
         "obligations. The same instruments apply to the removal of hazardous "
         "capability, where the difference between suppressed and removed "
         "knowledge is a safety question rather than a privacy one.")

    heading(doc, "5.4  Future Work Plan")
    body(doc,
         "The remaining work is organised into four tasks, in priority order.")
    subheading(doc, "T1  Separate targeted forgetting from general degradation")
    body(doc,
         "The utility control fell by 0.095 across the extrapolation sweep "
         "reported in Section 4.4.4, so the joint rise in the two axes is not "
         "yet attributable to forgetting alone. Two measures address this. A "
         "checkpoint will be selected with more headroom above the utility "
         "floor, so that the sweep begins from a less strained model. Each point "
         "will additionally be compared against a degradation control: a model "
         "perturbed by an equal-magnitude update in a direction unrelated to the "
         "forget set, which should move utility similarly while leaving both "
         "measurement axes near zero. The difference isolates the targeted "
         "component. This task is a precondition for interpreting the "
         "trajectory, and is therefore first.")
    subheading(doc, "T2  Replicate across methods and seeds")
    body(doc,
         "The negative preference optimisation and gradient difference arms will "
         "be run under the identical protocol, each with three random seeds. "
         "Negative preference optimisation is of particular interest because it "
         "was designed to avoid the collapse behaviour visible in Figure 10, so "
         "it should sustain a wider usable range of unlearning strength. This "
         "task converts a single trajectory into a set of independent "
         "observations, which is what permits a claim about the depth-breadth "
         "relationship to be made at all.")
    subheading(doc, "T3  Extend the breadth axis")
    body(doc,
         "The alias, entailment and multi-hop tiers will be authored and added. "
         "These probe the routes by which suppressed knowledge is most likely to "
         "remain accessible, and are the tiers at which a difference between "
         "deep and broad forgetting would be expected to appear most clearly. "
         "Each tier requires hand-written items validated against the reference "
         "models before use.")
    subheading(doc, "T4  Scale and consolidate")
    body(doc,
         "The complete protocol will be repeated at the 3B scale to establish "
         "whether the relationship is scale-dependent, and a parameter-efficient "
         "adaptation arm will be added as a second adaptation regime. Results, "
         "artefacts and documentation will be consolidated for the final "
         "report.")
    subheading(doc, "Points proposed for discussion with the mentor")
    for s in [
        "Priority of the degradation control (T1) against the replication (T2). "
        "T1 is required for the current trajectory to be interpretable, but T2 "
        "is required for any claim at all; the team proposes T1 first and seeks "
        "confirmation.",
        "Whether the negative preference optimisation arm should replace "
        "gradient ascent as the primary method for the headline result, given "
        "that gradient ascent's usable range is bounded by the collapse "
        "behaviour in Figure 10.",
        "Scope of the additional breadth tiers: whether to author all three "
        "(alias, entailment, multi-hop) or to prioritise the alias tier, which "
        "is the least ambiguous to construct and validate.",
        "Whether the 3B scale is a necessary component of the final report or a "
        "desirable extension, given the compute it requires relative to "
        "additional seeds at the 1B scale.",
        "Target venue and format for the final write-up, and whether the "
        "measurement package should be released publicly alongside it.",
        "Confirmation that the objectives as stated in Table 3 match the panel's "
        "approved wording.",
    ]:
        bullet(doc, s)


# ===========================================================================
# APPENDICES
# ===========================================================================
REFERENCES = [
    "P. Maini, Z. Feng, A. Schwarzschild, Z. C. Lipton, and J. Z. Kolter, "
    "\"TOFU: A task of fictitious unlearning for LLMs,\" in Proc. First Conf. "
    "Language Modeling (COLM), 2024.",

    "J. Lee, D. Kim, and J. Jo, \"Measuring the depth of LLM unlearning,\" "
    "arXiv preprint arXiv:2605.24614, 2026.",

    "W. Wang, M. Zhang, X. Ye, Z. Ren, P. Ren, and Z. Chen, \"UIPE: Enhancing "
    "LLM unlearning by removing knowledge related to forgetting targets,\" in "
    "Findings Assoc. Comput. Linguistics: EMNLP 2025, 2025, pp. 25212-25227.",

    "R. Zhang, L. Lin, Y. Bai, and S. Mei, \"Negative preference optimization: "
    "From catastrophic collapse to effective unlearning,\" in Proc. First Conf. "
    "Language Modeling (COLM), 2024.",

    "J. Jang, D. Yoon, S. Yang, S. Cha, M. Lee, L. Logeswaran, and M. Seo, "
    "\"Knowledge unlearning for mitigating privacy risks in language models,\" "
    "in Proc. 61st Annu. Meeting Assoc. Comput. Linguistics (ACL), 2023.",

    "R. Eldan and M. Russinovich, \"Who's Harry Potter? Approximate unlearning "
    "in LLMs,\" arXiv preprint arXiv:2310.02238, 2023.",

    "K. Meng, D. Bau, A. Andonian, and Y. Belinkov, \"Locating and editing "
    "factual associations in GPT,\" in Advances in Neural Information "
    "Processing Systems (NeurIPS), 2022.",

    "N. Li, A. Pan, A. Gopal, S. Yue, D. Berrios, A. Gatti, J. D. Li, "
    "A.-K. Dombrowski, et al., \"The WMDP benchmark: Measuring and reducing "
    "malicious use with unlearning,\" in Proc. 41st Int. Conf. Machine Learning "
    "(ICML), 2024.",

    "W. Shi, J. Lee, Y. Huang, S. Malladi, J. Zhao, A. Holtzman, D. Liu, "
    "L. Zettlemoyer, et al., \"MUSE: Machine unlearning six-way evaluation for "
    "language models,\" in Proc. Int. Conf. Learning Representations (ICLR), "
    "2025.",

    "Z. Jin, P. Cao, C. Wang, Z. He, H. Yuan, J. Li, Y. Chen, K. Liu, and "
    "J. Zhao, \"RWKU: Benchmarking real-world knowledge unlearning for large "
    "language models,\" in Advances in Neural Information Processing Systems "
    "(NeurIPS) Datasets and Benchmarks Track, 2024, pp. 98213-98263.",

    "A. Lynch, P. Guo, A. Ewart, S. Casper, and D. Hadfield-Menell, \"Eight "
    "methods to evaluate robust unlearning in LLMs,\" arXiv preprint "
    "arXiv:2402.16835, 2024.",

    "J. Łucki, B. Wei, Y. Huang, P. Henderson, F. Tramèr, and J. Rando, \"An "
    "adversarial perspective on machine unlearning for AI safety,\" "
    "Transactions on Machine Learning Research (TMLR), 2025.",

    "A. Deeb and F. Roger, \"Do unlearning methods remove information from "
    "language model weights?\" arXiv preprint arXiv:2410.08827, 2024.",

    "J. Doshi and A. C. Stickland, \"Does unlearning truly unlearn? A black box "
    "evaluation of LLM unlearning methods,\" arXiv preprint arXiv:2411.12103, "
    "2024.",

    "Y. Hong, L. Yu, H. Yang, S. Ravfogel, and M. Geva, \"Intrinsic test of "
    "unlearning using parametric knowledge traces,\" in Proc. 2025 Conf. "
    "Empirical Methods in Natural Language Processing (EMNLP), 2025.",

    "H. Lee, U. Hwang, H. Lim, and T. Kim, \"Does localization inform "
    "unlearning? A rigorous examination of local parameter attribution for "
    "knowledge unlearning in language models,\" in Proc. 2025 Conf. Empirical "
    "Methods in Natural Language Processing (EMNLP), 2025, pp. 21857-21869.",

    "G. Ilharco, M. T. Ribeiro, M. Wortsman, S. Gururangan, L. Schmidt, "
    "H. Hajishirzi, and A. Farhadi, \"Editing models with task arithmetic,\" in "
    "Proc. Int. Conf. Learning Representations (ICLR), 2023.",

    "S. Liu, Y. Yao, J. Jia, S. Casper, N. Baracaldo, P. Hase, Y. Yao, "
    "C. Y. Liu, et al., \"Rethinking machine unlearning for large language "
    "models,\" Nature Machine Intelligence, 2025.",

    "F. Barez, T. Fu, A. Prabhu, S. Casper, A. Sanyal, A. Bibi, A. O'Gara, "
    "R. Kirk, et al., \"Open problems in machine unlearning for AI safety,\" "
    "arXiv preprint arXiv:2501.04952, 2025.",

    "P. Thaker, S. Hu, N. Kale, et al., \"Position: LLM unlearning benchmarks "
    "are weak measures of progress,\" in Proc. IEEE Conf. Secure and "
    "Trustworthy Machine Learning (SaTML), 2025.",

    "E. J. Hu, Y. Shen, P. Wallis, Z. Allen-Zhu, Y. Li, S. Wang, L. Wang, and "
    "W. Chen, \"LoRA: Low-rank adaptation of large language models,\" in Proc. "
    "Int. Conf. Learning Representations (ICLR), 2022.",

    "K. Meng, A. Sen Sharma, A. Andonian, Y. Belinkov, and D. Bau, "
    "\"Mass-editing memory in a transformer,\" in Proc. Int. Conf. Learning "
    "Representations (ICLR), 2023.",

    "C. Fan, J. Liu, et al., \"Simplicity prevails: Rethinking negative "
    "preference optimization for LLM unlearning,\" arXiv preprint "
    "arXiv:2410.07163, 2024.",

    "C. Y. Liu, Y. Wang, J. Flanigan, and Y. Liu, \"Large language model "
    "unlearning via embedding-corrupted prompts,\" in Advances in Neural "
    "Information Processing Systems (NeurIPS), 2024.",

    "B. W. Lee, A. Foote, A. Infanger, L. Shor, H. Kamath, J. Goldman-Wetzler, "
    "B. Woodworth, A. Cloud, et al., \"Distillation robustifies unlearning,\" in "
    "Advances in Neural Information Processing Systems (NeurIPS), 2025.",

    "X. Wang, Z. Li, B. Wang, et al., \"Model unlearning via sparse autoencoder "
    "subspace guided projections,\" in Proc. 2025 Conf. Empirical Methods in "
    "Natural Language Processing (EMNLP), 2025.",
]


def _appendices(doc):
    doc.add_page_break()
    para(doc, "APPENDIX A: REFERENCES", size=CHAP_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    rule(doc)
    para(doc, "", space_after=8)
    for i, ref in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.0        # single spaced
        p.paragraph_format.space_after = Pt(10)      # double space between
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        _set_font(p.add_run(f"[{i}]\t"), BODY_PT)
        _set_font(p.add_run(ref), BODY_PT)

    doc.add_page_break()
    para(doc, "APPENDIX B: PLAGIARISM REPORT", size=CHAP_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    rule(doc)
    para(doc, "", space_after=8)
    body(doc, "The similarity report generated by the institute's plagiarism "
              "detection service is to be attached here.", italic=True,
         color=RED)


if __name__ == "__main__":
    build()
